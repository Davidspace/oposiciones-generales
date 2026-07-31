from __future__ import annotations

import json
import re
import sys
from collections import Counter
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_materials as bm


DATA_DIR = ROOT / "_trabajo" / "investigacion" / "analisis_examenes"
PROGRAM_PATH = ROOT / "_trabajo" / "investigacion" / "programa_maestro.json"
OUT_DIR = ROOT / "08. Análisis de exámenes y preguntas"
QA_DIR = ROOT / "_trabajo" / "qa" / "graficos_analisis"
SOURCE_MD = ROOT / "_trabajo" / "contenidos_auxiliares" / "informe_tendencias_examenes.md"
OUT_DOCX = OUT_DIR / "Informe de tendencias de exámenes oficiales.docx"

NAVY = "#17365D"
BLUE = "#2F5597"
MID_BLUE = "#5B9BD5"
LIGHT_BLUE = "#D9EAF7"
GRAY = "#64748B"
FONT_DIR = Path("C:/Windows/Fonts")


def load_json(name: str):
    return json.loads((DATA_DIR / name).read_text(encoding="utf-8"))


def norm_key(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip()).casefold()


def shorten(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip(" ,.;:") + "…"


def rows_to_md(headers: list[str], rows: list[list[object]]) -> str:
    def safe(value: object) -> str:
        return str(value).replace("|", "/").replace("\n", " ")

    lines = [
        "| " + " | ".join(safe(x) for x in headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    lines.extend("| " + " | ".join(safe(x) for x in row) + " |" for row in rows)
    return "\n".join(lines)


def aggregate_casefold(rows: list[dict], field: str) -> list[tuple[str, int]]:
    labels: dict[str, str] = {}
    counter: Counter[str] = Counter()
    for row in rows:
        value = str(row.get(field, "")).strip()
        if not value:
            continue
        key = norm_key(value)
        counter[key] += 1
        existing = labels.get(key)
        if existing is None or (value[:1].isupper() and not existing[:1].isupper()):
            labels[key] = value
    return [(labels[key], count) for key, count in counter.most_common()]


def chart_font(size: int, *, bold: bool = False):
    filename = "arialbd.ttf" if bold else "arial.ttf"
    try:
        return ImageFont.truetype(str(FONT_DIR / filename), size=size)
    except OSError:
        return ImageFont.load_default()


def draw_horizontal_chart(
    labels: list[str],
    values: list[int],
    *,
    title: str,
    path: Path,
    colors: list[str] | None = None,
    value_labels: list[str] | None = None,
) -> Path:
    width = 2200
    row_height = 92
    top = 180
    bottom = 95
    height = top + bottom + row_height * len(labels)
    label_right = 920
    chart_left = 970
    chart_right = 2010
    max_value = max(values) if values else 1
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = chart_font(48, bold=True)
    label_font = chart_font(30)
    value_font = chart_font(30, bold=True)
    small_font = chart_font(25)
    draw.text((55, 42), title, fill=NAVY, font=title_font)
    draw.line((55, 122, 2145, 122), fill=LIGHT_BLUE, width=4)

    grid_steps = 5
    for step in range(grid_steps + 1):
        x = chart_left + round((chart_right - chart_left) * step / grid_steps)
        draw.line((x, top - 12, x, height - bottom + 6), fill="#E2E8F0", width=2)
        tick = round(max_value * step / grid_steps)
        tick_text = str(tick)
        bbox = draw.textbbox((0, 0), tick_text, font=small_font)
        draw.text(
            (x - (bbox[2] - bbox[0]) / 2, height - bottom + 22),
            tick_text,
            fill=GRAY,
            font=small_font,
        )

    palette = colors or [BLUE] * len(labels)
    labels_out = value_labels or [str(value) for value in values]
    for index, (label, value, color, value_label) in enumerate(
        zip(labels, values, palette, labels_out)
    ):
        y = top + index * row_height
        label_box = draw.textbbox((0, 0), label, font=label_font)
        draw.text(
            (label_right - (label_box[2] - label_box[0]), y + 15),
            label,
            fill="#334155",
            font=label_font,
        )
        bar_width = max(3, round((chart_right - chart_left) * value / max_value))
        draw.rounded_rectangle(
            (chart_left, y + 10, chart_left + bar_width, y + 58),
            radius=13,
            fill=color,
        )
        draw.text(
            (chart_left + bar_width + 18, y + 14),
            value_label,
            fill=NAVY,
            font=value_font,
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, "PNG", optimize=True)
    return path


def save_top_theme_chart(freq: list[dict]) -> Path:
    top = sorted(freq, key=lambda row: row["preguntas"], reverse=True)[:12]
    labels = [
        f"{row['tema_id']} · {shorten(row['titulo'].split('.')[0], 39)}"
        for row in top
    ]
    values = [row["preguntas"] for row in top]
    return draw_horizontal_chart(
        labels,
        values,
        title="Temas con mayor presencia en cinco sesiones canónicas",
        path=QA_DIR / "01_frecuencia_temas.png",
        colors=[BLUE if value >= 30 else MID_BLUE for value in values],
    )


def save_block_chart(canonical: list[dict]) -> Path:
    counts = Counter(row["bloque"] for row in canonical)
    values = [
        counts.get("Temario general", 0),
        counts.get("Temario específico Seguridad Social", 0),
    ]
    total = sum(values)
    return draw_horizontal_chart(
        ["Temario general", "Seguridad Social"],
        values,
        title="Peso observado por bloque",
        path=QA_DIR / "02_peso_bloques.png",
        colors=[MID_BLUE, BLUE],
        value_labels=[f"{value} · {value / total:.1%}" for value in values],
    )


def save_type_chart(canonical: list[dict]) -> Path:
    top = Counter(row["tipo_pregunta"] for row in canonical).most_common(10)
    return draw_horizontal_chart(
        [label for label, _ in top],
        [value for _, value in top],
        title="Formatos de pregunta más repetidos",
        path=QA_DIR / "03_tipos_pregunta.png",
        colors=[BLUE] * len(top),
    )


def add_figure(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(16.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    bm.set_run_font(r, size=8.5, italic=True, color=bm.MID_GRAY)


def build_markdown(
    questions: list[dict],
    canonical: list[dict],
    freq: list[dict],
    inventory: list[dict],
    summary: dict,
) -> tuple[str, str, str, str]:
    counts_block = Counter(row["bloque"] for row in canonical)
    counts_year = Counter(row["año"] for row in canonical)
    counts_character = Counter(row["caracter"] for row in canonical)
    counts_reserve = Counter(row["reserva"] for row in canonical)
    counts_result_all = Counter(row["resultado_final"] for row in questions)
    counts_result_canonical = Counter(row["resultado_final"] for row in canonical)

    top_freq = sorted(freq, key=lambda row: row["preguntas"], reverse=True)
    top_rows = [
        [
            row["tema_id"],
            shorten(row["titulo"], 78),
            row["preguntas"],
            row["practicas"],
            row["reservas"],
            ", ".join(map(str, row["años"])),
        ]
        for row in top_freq[:15]
    ]
    annual_rows = [
        [year, counts_year[year], f"{counts_year[year] / len(canonical):.1%}"]
        for year in sorted(counts_year)
    ]
    inventory_rows = [
        [
            row["convocatoria"],
            row["fecha_examen"],
            row["modelo"],
            row["preguntas"],
            row["reservas"],
            row["plantilla_definitiva"],
            shorten(row["observaciones"], 76),
        ]
        for row in inventory
    ]

    intro = f"""
# Resumen ejecutivo

Este informe estudia el corpus oficial disponible para el Cuerpo Administrativo de la Administración de la Seguridad Social, acceso libre, con corte documental **30 de julio de 2026**. La base conserva **{summary['preguntas_registradas']} registros** procedentes de seis variantes publicadas y utiliza **{summary['preguntas_canonicas_para_frecuencia']} preguntas canónicas** para calcular frecuencias sin duplicar modelos equivalentes.

| Indicador | Resultado |
| --- | --- |
| Sesiones canónicas | {summary['sesiones_canonicas']} |
| Preguntas canónicas | {len(canonical)} |
| Temario general | {counts_block['Temario general']} ({counts_block['Temario general'] / len(canonical):.1%}) |
| Temario específico de Seguridad Social | {counts_block['Temario específico Seguridad Social']} ({counts_block['Temario específico Seguridad Social'] / len(canonical):.1%}) |
| Preguntas teóricas / prácticas | {counts_character['Teórico']} / {counts_character['Práctico']} |
| Preguntas ordinarias / de reserva | {counts_reserve['No']} / {counts_reserve['Sí']} |
| Anuladas en las seis variantes / canónicas | {counts_result_all['Anulada']} / {counts_result_canonical['Anulada']} |
| Respuestas modificadas | {counts_result_all['Modificada']} |
| Registros de baja confianza | {summary['baja_confianza']} |

La señal más clara es la concentración en el bloque específico: casi tres cuartas partes de las preguntas observadas se relacionan con Seguridad Social. Dentro de ese bloque destacan incapacidad, jubilación, afiliación, cotización y recaudación. En el bloque general el patrón es más disperso y predominan la literalidad normativa, el procedimiento administrativo y el empleo público.

## Alcance y criterio de cómputo

La base registra las preguntas 1–70 y sus tres reservas de la primera parte, así como las 15 preguntas del supuesto y sus tres reservas: **91 registros por variante**. Los modelos B permanecen archivados como PDF oficial, pero no se vuelcan como una segunda batería cuando solo permutan orden u opciones. En 2026 se conservaron Rojo A y Verde A para trazabilidad; Verde A se enlaza con su equivalente canónico y se excluye de frecuencia.

{rows_to_md(['Conv.', 'Fecha', 'Modelo', 'Preguntas', 'Reservas', 'Definitiva', 'Observación'], inventory_rows)}

### Frecuencia temporal

{rows_to_md(['Año', 'Preguntas canónicas', 'Peso'], annual_rows)}

El volumen anual refleja el número de sesiones disponibles, no una intensificación automática del temario: 2025 reúne tres sesiones canónicas y por eso concentra más registros.

## Metodología de construcción y control

1. Descarga y archivo de cuestionarios, plantillas provisionales, plantillas definitivas y resoluciones oficiales.
2. Extracción de cada enunciado y sus cuatro opciones mediante OCR, con reparación documentada de tres opciones ausentes en la capa de texto.
3. Enlace de respuestas provisionales y definitivas; identificación separada de anulaciones y cambios.
4. Clasificación por tema oficial, apartado, norma, artículo, tipo, dificultad y carácter teórico o práctico.
5. Deduplicación de variantes equivalentes antes de calcular frecuencia.
6. Controles automáticos: 91 preguntas por variante, cuatro alternativas por pregunta, ausencia de campos de opción vacíos y cero clasificaciones de baja confianza.

Las frecuencias son descriptivas. Indican lo preguntado en el corpus conservado, pero no convierten una tendencia en predicción ni sustituyen el estudio completo del programa oficial.
"""

    themes_section = f"""
# Frecuencia por temas

{rows_to_md(['Tema', 'Denominación abreviada', 'Total', 'Prácticas', 'Reservas', 'Años'], top_rows)}

Los cinco primeros temas por volumen son **{', '.join(row['tema_id'] for row in top_freq[:5])}**. Su concentración responde, sobre todo, a preguntas de aplicación, plazos, requisitos, bases reguladoras, porcentajes y secuencias de gestión.
"""

    norm_counts = aggregate_casefold(canonical, "norma")
    article_counts = aggregate_casefold(canonical, "articulo_precepto")
    type_counts = Counter(row["tipo_pregunta"] for row in canonical).most_common(15)
    norm_rows = [[name, count] for name, count in norm_counts[:15]]
    article_rows = [[article, count] for article, count in article_counts[:20]]
    type_rows = [[name, count, f"{count / len(canonical):.1%}"] for name, count in type_counts]

    patterns = f"""
# Patrones de pregunta

{rows_to_md(['Tipo', 'Preguntas', 'Peso'], type_rows)}

La literalidad de artículos es el formato dominante, pero no debe confundirse con memorización aislada. En el supuesto práctico se combinan reglas de competencia, plazos, requisitos y operaciones numéricas. La preparación más eficiente une la literalidad con microcasos: hecho causante, órgano competente, plazo aplicable, porcentaje y consecuencia jurídica.

## Normas más citadas

{rows_to_md(['Norma identificada', 'Preguntas'], norm_rows)}

## Artículos y preceptos más repetidos

{rows_to_md(['Artículo o precepto', 'Preguntas'], article_rows)}

Los nombres de normas se han normalizado para el recuento sin alterar el texto trazable guardado en cada registro. Un artículo repetido de normas distintas no se agrega como si fuera el mismo precepto cuando la norma puede identificarse.
"""

    practical_rows = [
        [
            row["tema_id"],
            shorten(row["titulo"], 72),
            row["practicas"],
            row["preguntas"],
            f"{row['practicas'] / row['preguntas']:.1%}" if row["preguntas"] else "0,0 %",
        ]
        for row in sorted(freq, key=lambda row: row["practicas"], reverse=True)
        if row["practicas"]
    ][:15]
    reserve_rows = [
        [row["tema_id"], shorten(row["titulo"], 74), row["reservas"], row["preguntas"]]
        for row in sorted(freq, key=lambda row: row["reservas"], reverse=True)
        if row["reservas"]
    ][:15]
    changes = [
        row
        for row in questions
        if row["resultado_final"] in {"Anulada", "Modificada"}
    ]
    change_rows = [
        [
            row["id_pregunta"],
            row["tema_id"],
            row["resultado_final"],
            row["respuesta_provisional"] or "—",
            row["respuesta_definitiva"] or "—",
            shorten(row["concepto_principal"], 58),
        ]
        for row in changes
    ]

    application = f"""
# Supuesto práctico, reservas y cambios de plantilla

## Temas con mayor presencia práctica

{rows_to_md(['Tema', 'Denominación abreviada', 'Prácticas', 'Total', '% práctico'], practical_rows)}

## Temas con más preguntas de reserva

{rows_to_md(['Tema', 'Denominación abreviada', 'Reservas', 'Total'], reserve_rows)}

## Anulaciones y modificación de respuesta

{rows_to_md(['ID', 'Tema', 'Resultado', 'Provisional', 'Definitiva', 'Concepto'], change_rows)}

Las anulaciones deben estudiarse como incidencias documentales, no como frecuencia positiva de una regla. La hoja de cálculo permite filtrar cada caso y regresar al cuestionario y a la plantilla oficial correspondiente.
"""

    full_rows = [
        [
            row["tema_id"],
            row["bloque"].replace("Temario específico Seguridad Social", "Seguridad Social").replace("Temario general", "General"),
            shorten(row["titulo"], 68),
            row["preguntas"],
            row["practicas"],
            row["reservas"],
            row["anuladas"],
            ", ".join(map(str, row["años"])) if row["años"] else "—",
        ]
        for row in sorted(freq, key=lambda row: (row["tema_id"][0], row["numero"]))
    ]
    conclusions = f"""
# Mapa completo de los 36 temas

{rows_to_md(['ID', 'Bloque', 'Denominación abreviada', 'Total', 'Práct.', 'Res.', 'Anul.', 'Años'], full_rows)}

# Conclusiones para la preparación

1. **Cobertura completa antes que selección extrema.** Los 36 temas aparecen en el material y una frecuencia baja no equivale a exclusión futura.
2. **Prioridad alta al núcleo específico.** S03–S12 reúnen la mayor densidad de preguntas, cálculos y microcasos. Conviene entrenarlos con revisión espaciada y supuestos cronometrados.
3. **Literalidad con contexto.** Constitución, Ley 39/2015, Ley 40/2015, TREBEP y normativa de Seguridad Social exigen recordar artículo, órgano, plazo y excepción.
4. **Reservas incluidas siempre.** La resolución del examen debe abarcar las preguntas de reserva porque pueden sustituir a anuladas.
5. **Control de errores con penalización.** La convocatoria vigente descuenta un cuarto del valor de una respuesta correcta por cada fallo; la estrategia de intento debe basarse en descarte razonado.
6. **Repaso orientado a datos.** Edades, porcentajes, bases, duraciones y plazos deben separarse en fichas de cifras para reducir confusiones próximas.
7. **Revisión normativa fechada.** Los resultados describen preguntas anteriores. Para responder un examen posterior debe verificarse siempre la norma vigente en la fecha de corte que señale el órgano de selección.

## Limitaciones

- El corpus depende de la documentación oficial publicada y conservada en el expediente.
- El OCR puede mantener erratas tipográficas de los cuestionarios; las cuatro opciones y la correspondencia de plantilla fueron controladas.
- Los modelos equivalentes no se suman dos veces; por ello los recuentos pueden ser inferiores al total de PDF archivados.
- Una pregunta transversal se asigna al tema que mejor representa su regla decisiva. El texto, la norma y el artículo permanecen disponibles para reauditar esa decisión.
- La frecuencia no mide por sí sola dificultad, valor pedagógico ni probabilidad futura.

# Fuentes oficiales

- Convocatoria vigente, Resolución de 22 de diciembre de 2025: https://www.boe.es/diario_boe/txt.php?id=BOE-A-2025-27158
- Corrección de errores de 4 de marzo de 2026: https://www.boe.es/diario_boe/txt.php?id=BOE-A-2026-5351
- Portal oficial del proceso selectivo y documentación de examen: https://www.seg-social.es/wps/portal/wss/internet/InformacionUtil/9950/88beb45d-ca83-4bab-8162-f94ef2894562/e0c83707-3358-47a9-94d7-508221dad233
- Oferta de Empleo Público de 2026, como referencia separada de la convocatoria vigente: https://boe.es/diario_boe/txt.php?id=BOE-A-2026-9946

La base detallada, los filtros, el inventario y los recuentos reproducibles se entregan en **Análisis de exámenes oficiales.xlsx**.
"""
    return intro.strip(), themes_section.strip(), patterns.strip(), (application + "\n\n" + conclusions).strip()


def build_report() -> dict:
    questions = load_json("preguntas_oficiales.json")
    freq = load_json("mapa_frecuencia.json")
    inventory = load_json("inventario_examenes.json")
    summary = load_json("resumen_analisis.json")
    canonical = [row for row in questions if row["computar_en_frecuencia"] == "Sí"]

    if len(canonical) != summary["preguntas_canonicas_para_frecuencia"]:
        raise ValueError("El total canónico no coincide con el resumen de análisis.")
    if any(not all(row.get(f"opcion_{letter}") for letter in "abcd") for row in questions):
        raise ValueError("Hay preguntas sin cuatro opciones.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    SOURCE_MD.parent.mkdir(parents=True, exist_ok=True)
    charts = [
        save_top_theme_chart(freq),
        save_block_chart(canonical),
        save_type_chart(canonical),
    ]
    sections = build_markdown(questions, canonical, freq, inventory, summary)
    SOURCE_MD.write_text("\n\n".join(sections) + "\n", encoding="utf-8")

    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(
        doc, "Informe de tendencias · Exámenes oficiales · Actualizado 30-07-2026"
    )
    bm.add_cover(
        doc,
        kind="informe de análisis",
        number=None,
        title="Informe de tendencias de exámenes oficiales",
        block="Cuerpo Administrativo de la Administración de la Seguridad Social",
    )

    bm.add_markdown(doc, sections[0])
    add_figure(
        doc,
        charts[1],
        "Figura 1. Distribución del corpus canónico entre temario general y Seguridad Social.",
    )
    bm.add_markdown(doc, sections[1])
    add_figure(
        doc,
        charts[0],
        "Figura 2. Doce temas con mayor frecuencia observada.",
    )
    bm.add_markdown(doc, sections[2])
    add_figure(
        doc,
        charts[2],
        "Figura 3. Diez formatos de pregunta más repetidos.",
    )
    bm.add_markdown(doc, sections[3])

    doc.core_properties.title = "Informe de tendencias de exámenes oficiales"
    doc.core_properties.subject = (
        "Cuerpo Administrativo de la Administración de la Seguridad Social, acceso libre"
    )
    doc.core_properties.author = "ACADEMIA LORMAN"
    doc.core_properties.keywords = "Seguridad Social; oposición; C1; exámenes; frecuencia"
    doc.save(OUT_DOCX)

    return {
        "fecha_corte": "2026-07-30",
        "docx": str(OUT_DOCX.relative_to(ROOT)),
        "fuente_markdown": str(SOURCE_MD.relative_to(ROOT)),
        "graficos": [str(path.relative_to(ROOT)) for path in charts],
        "preguntas_registradas": len(questions),
        "preguntas_canonicas": len(canonical),
        "temas": len(freq),
    }


if __name__ == "__main__":
    print(json.dumps(build_report(), ensure_ascii=False, indent=2))

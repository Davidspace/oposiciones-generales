from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
QA_DIR = ROOT / "_trabajo" / "qa"
GENERATION = QA_DIR / "informe_generacion.json"
SIM_AUDIT = QA_DIR / "auditoria_simulacros.json"
CONTROL = QA_DIR / "control_final_entregables.json"
DUPLICATES = QA_DIR / "auditoria_duplicados_preguntas.json"
THEME_PAGES = QA_DIR / "auditoria_paginas_temarios.json"
EXPECTED_COPYRIGHT = "© 2026 ACADEMIA LORMAN"
OLD_MARKERS = ("© 2026 Alba",)
LEGACY_DATE_MARKERS = ("29 de julio de 2026", "29-07-2026", "2026-07-29")
INTERNAL_MARKERS = (
    "generado por ia",
    "generada por ia",
    "chatgpt",
    "openai",
    "codex",
    "modelo de lenguaje",
)


def all_generated_docx() -> list[Path]:
    paths: list[Path] = []
    for folder in (
        "01. Información oficial de la oposición",
        "02. Temas redactados",
        "03. Test por temas",
        "04. Simulacros",
        "05. Normativa",
        "06. Fuentes y control de actualización",
        "07. Informes de revisión",
        "08. Análisis de exámenes y preguntas",
    ):
        paths.extend((ROOT / folder).rglob("*.docx"))
    return sorted(paths)


def package_text(path: Path) -> tuple[str, list[str], str]:
    names: list[str] = []
    chunks: list[str] = []
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if name.endswith(".xml"):
                names.append(name)
                chunks.append(archive.read(name).decode("utf-8", errors="replace"))
    raw = "\n".join(chunks)
    visible = re.sub(r"<[^>]+>", " ", raw)
    visible = (
        visible.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
    )
    return re.sub(r"\s+", " ", visible), names, raw


def audit_docx(path: Path) -> dict:
    relative = str(path.relative_to(ROOT))
    result = {
        "path": relative,
        "integrity": False,
        "copyright": False,
        "current_date": False,
        "a4": False,
        "arial": False,
        "branded_author": False,
        "different_first_page": False,
        "old_markers": [],
        "legacy_date_mentions": [],
        "internal_markers": [],
        "paragraphs": 0,
        "tables": 0,
        "bytes": path.stat().st_size,
    }
    try:
        text, names, raw_xml = package_text(path)
        document = Document(path)
    except (BadZipFile, KeyError, ValueError, OSError) as error:
        result["error"] = str(error)
        return result

    result["integrity"] = True
    result["copyright"] = EXPECTED_COPYRIGHT in text
    result["current_date"] = (
        "30 de julio de 2026" in text
        or "30-07-2026" in text
        or "2026-07-30" in text
    )
    result["arial"] = "Arial" in raw_xml
    result["branded_author"] = (
        document.core_properties.author == "ACADEMIA LORMAN"
        and document.core_properties.last_modified_by == "ACADEMIA LORMAN"
    )
    result["paragraphs"] = len(document.paragraphs)
    result["tables"] = len(document.tables)
    result["different_first_page"] = all(
        section.different_first_page_header_footer for section in document.sections
    )
    result["a4"] = all(
        abs(section.page_width.mm - 210) < 1
        and abs(section.page_height.mm - 297) < 1
        for section in document.sections
    )
    result["old_markers"] = [marker for marker in OLD_MARKERS if marker in text]
    result["legacy_date_mentions"] = [
        marker for marker in LEGACY_DATE_MARKERS if marker in text
    ]
    lower_text = text.casefold()
    result["internal_markers"] = [
        marker for marker in INTERNAL_MARKERS if marker in lower_text
    ]
    result["footer_parts"] = sum(name.startswith("word/footer") for name in names)
    return result


def audit_pdf(path: Path) -> dict:
    result = {
        "path": str(path.relative_to(ROOT)),
        "integrity": False,
        "pages": 0,
        "bytes": path.stat().st_size,
        "encrypted": False,
        "empty_text_pages": [],
    }
    try:
        reader = PdfReader(str(path))
        result["encrypted"] = reader.is_encrypted
        result["pages"] = len(reader.pages)
        for index, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if len(text) < 20:
                result["empty_text_pages"].append(index)
        result["integrity"] = result["pages"] > 0 and not result["encrypted"]
    except Exception as error:
        result["error"] = str(error)
    return result


def main() -> None:
    generation = json.loads(GENERATION.read_text(encoding="utf-8"))
    simulation = json.loads(SIM_AUDIT.read_text(encoding="utf-8"))
    control = json.loads(CONTROL.read_text(encoding="utf-8"))
    duplicates = json.loads(DUPLICATES.read_text(encoding="utf-8"))
    theme_pages = json.loads(THEME_PAGES.read_text(encoding="utf-8"))

    docx_paths = all_generated_docx()
    docx_audit = [audit_docx(path) for path in docx_paths]
    key_pdf_paths = [
        ROOT / "01. Información oficial de la oposición" / "Guía completa de la oposición.pdf",
        ROOT
        / "08. Análisis de exámenes y preguntas"
        / "Informe de tendencias de exámenes oficiales.pdf",
    ]
    pdf_audit = [audit_pdf(path) for path in key_pdf_paths]

    theme_words = {
        key: value["word_count"] for key, value in generation["themes"].items()
    }
    invalid_tests = [
        key for key, value in generation["tests"].items() if not value["valid"]
    ]
    failure_fields = (
        "integrity",
        "copyright",
        "current_date",
        "a4",
        "arial",
        "branded_author",
        "different_first_page",
    )
    docx_failures = []
    for item in docx_audit:
        failures = [field for field in failure_fields if not item.get(field)]
        failures.extend(f"old:{marker}" for marker in item["old_markers"])
        failures.extend(f"internal:{marker}" for marker in item["internal_markers"])
        if failures:
            docx_failures.append({"path": item["path"], "failures": failures})

    counts = Counter()
    for path in docx_paths:
        relative = path.relative_to(ROOT)
        counts[relative.parts[0]] += 1

    report = {
        "fecha_corte": "2026-07-30",
        "counts_docx_by_folder": dict(sorted(counts.items())),
        "docx_total": len(docx_paths),
        "docx_failures": docx_failures,
        "themes": {
            "count": len(generation["themes"]),
            "minimum_words": min(theme_words.values()),
            "maximum_words": max(theme_words.values()),
            "below_5500": [
                key for key, value in theme_words.items() if value < 5500
            ],
        },
        "tests": {
            "count": len(generation["tests"]),
            "invalid": invalid_tests,
        },
        "simulations": simulation["audit"],
        "question_duplicates": {
            "questions": duplicates["questions"],
            "duplicate_exact_stems": duplicates["duplicate_exact_stems"],
            "near_duplicate_pairs_090": len(
                duplicates["near_duplicate_pairs_090"]
            ),
            "near_pairs": duplicates["near_duplicate_pairs_090"],
        },
        "theme_page_audit": {
            "count": theme_pages["count"],
            "minimum_pages": theme_pages["minimum_pages"],
            "maximum_pages": theme_pages["maximum_pages"],
            "below_15_pages": theme_pages["below_15_pages"],
        },
        "control": control,
        "key_pdfs": pdf_audit,
        "analysis_workbook": {
            "exists": (
                ROOT
                / "08. Análisis de exámenes y preguntas"
                / "Análisis de exámenes oficiales.xlsx"
            ).exists(),
            "inspect_log": str(
                (
                    QA_DIR / "analisis_examenes_xlsx_inspect.ndjson"
                ).relative_to(ROOT)
            ),
        },
        "pass": (
            not docx_failures
            and not invalid_tests
            and not [key for key, value in theme_words.items() if value < 5500]
            and simulation["audit"]["duplicate_rendered_stems"] == 0
            and simulation["audit"]["invalid_docx"] == []
            and duplicates["duplicate_exact_stems"] == 0
            and theme_pages["below_15_pages"] == []
            and all(item["integrity"] for item in pdf_audit)
        ),
    }
    output = QA_DIR / "auditoria_final_paquete.json"
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "pass": report["pass"],
                "docx_total": report["docx_total"],
                "docx_failures": len(docx_failures),
                "themes": report["themes"],
                "tests": report["tests"],
                "simulations": {
                    key: simulation["audit"][key]
                    for key in (
                        "simulations",
                        "student_versions",
                        "answer_templates",
                        "teacher_versions",
                        "questions_total_including_reserves",
                        "unique_rendered_stems",
                        "duplicate_rendered_stems",
                    )
                },
                "key_pdfs": pdf_audit,
                "report": str(output.relative_to(ROOT)),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()

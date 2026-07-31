from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_materials as bm  # noqa: E402


ROOT = SCRIPT_DIR.parents[1]
AUX_CONTENT_DIR = ROOT / "_trabajo" / "contenidos_auxiliares"
OFFICIAL_ROOT = ROOT / "01. Información oficial de la oposición"
THEME_ROOT = ROOT / "02. Temas redactados"
TEST_ROOT = ROOT / "03. Test por temas"
SIM_ROOT = ROOT / "04. Simulacros"
NORM_ROOT = ROOT / "05. Normativa"
SOURCE_ROOT = ROOT / "06. Fuentes y control de actualización"
REPORT_ROOT = ROOT / "07. Informes de revisión"
ANALYSIS_ROOT = ROOT / "08. Análisis de exámenes y preguntas"
QA_ROOT = ROOT / "_trabajo" / "qa"
PROGRAM_PATH = ROOT / "_trabajo" / "investigacion" / "programa_maestro.json"

UPDATE_DATE_TEXT = "30 de julio de 2026"
UPDATE_DATE_NUMERIC = "30-07-2026"


SUPPORT_SOURCES = {
    "guia": (
        AUX_CONTENT_DIR / "guia_completa.md",
        OFFICIAL_ROOT / "Guía completa de la oposición.docx",
        "GUÍA COMPLETA DE LA OPOSICIÓN",
        "Convocatoria, prueba, planificación y uso del material",
        "Guía · Actualizada 30-07-2026",
    ),
    "metodologia": (
        AUX_CONTENT_DIR / "metodologia_estudio.md",
        OFFICIAL_ROOT / "Metodología de estudio y planificación.docx",
        "METODOLOGÍA DE ESTUDIO",
        "Planificación, repasos, test y estrategia de examen",
        "Metodología · Actualizada 30-07-2026",
    ),
    "revision_juridica": (
        AUX_CONTENT_DIR / "informe_revision_juridica.md",
        REPORT_ROOT / "Informe de revisión jurídica y actualización normativa.docx",
        "INFORME DE REVISIÓN JURÍDICA",
        "Criterios, fecha de corte, alertas de vigencia y fuentes",
        "Revisión jurídica · 30-07-2026",
    ),
    "revision_editorial": (
        AUX_CONTENT_DIR / "informe_revision_editorial.md",
        REPORT_ROOT / "Informe de revisión editorial y técnica.docx",
        "INFORME DE REVISIÓN EDITORIAL",
        "Estructura, legibilidad, maquetación y controles automáticos",
        "Revisión editorial · 30-07-2026",
    ),
    "incidencias": (
        AUX_CONTENT_DIR / "informe_incidencias.md",
        REPORT_ROOT / "Informe de incidencias y decisiones editoriales.docx",
        "INFORME DE INCIDENCIAS",
        "Decisiones documentadas, limitaciones y medidas adoptadas",
        "Incidencias · 30-07-2026",
    ),
}


def ensure_directories() -> None:
    for path in (
        AUX_CONTENT_DIR,
        OFFICIAL_ROOT,
        SOURCE_ROOT,
        REPORT_ROOT,
        QA_ROOT,
    ):
        path.mkdir(parents=True, exist_ok=True)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(text))
    bm.set_run_font(r, size=8.5, bold=bold, color=color or bm.DARK_GRAY)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_grid_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths is None:
        widths = [round(bm.CONTENT_WIDTH_DXA / len(headers))] * len(headers)
        widths[-1] += bm.CONTENT_WIDTH_DXA - sum(widths)
    bm.set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_header(header)
    for idx, value in enumerate(headers):
        bm.shade_cell(header.cells[idx], bm.PALE_BLUE)
        set_cell_text(header.cells[idx], value, bold=True, color=bm.NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_support_cover(doc: Document, eyebrow: str, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CUERPO ADMINISTRATIVO DE LA ADMINISTRACIÓN DE LA SEGURIDAD SOCIAL")
    bm.set_run_font(r, size=9, bold=True, color=bm.MID_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(eyebrow)
    bm.set_run_font(r, size=13, bold=True, color=bm.BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    bm.set_run_font(r, size=21, bold=True, color=bm.NAVY)
    bm.add_bottom_border(p, color=bm.MID_BLUE, size="12")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    bm.set_run_font(r, size=11, italic=True, color=bm.DARK_GRAY)

    metadata = [
        ("Acceso", "Sistema general de acceso libre · Subgrupo C1"),
        ("Convocatoria base", "Resolución de 22 de diciembre de 2025 (BOE de 31-12-2025)"),
        ("Corrección", "BOE de 7 de marzo de 2026"),
        ("Fecha de corte del material", UPDATE_DATE_TEXT),
        ("Titular", "ACADEMIA LORMAN"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        bm.set_run_font(r1, size=9.5, bold=True, color=bm.DARK_GRAY)
        r2 = p.add_run(value)
        bm.set_run_font(r2, size=9.5, color=bm.DARK_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("AVISO DE PROPIEDAD INTELECTUAL")
    bm.set_run_font(r, size=8.5, bold=True, color=bm.BLUE)
    for idx, legal in enumerate(bm.LEGAL_PARAGRAPHS):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        bm.add_paragraph_shading(p, bm.LIGHT_GRAY)
        r = p.add_run(legal)
        bm.set_run_font(r, size=8, bold=(idx == 0), color=bm.DARK_GRAY)
    doc.add_page_break()


def build_markdown_document(key: str) -> dict:
    source, output, eyebrow, title, running = SUPPORT_SOURCES[key]
    if not source.exists():
        return {"key": key, "status": "omitted", "reason": f"No existe {source.name}"}
    markdown = source.read_text(encoding="utf-8")
    if "{{PROGRAMA_MAESTRO}}" in markdown:
        program = json.loads(PROGRAM_PATH.read_text(encoding="utf-8"))
        program_lines: list[str] = []
        for block in ("Temario general", "Temario de Seguridad Social"):
            program_lines.append(f"### {block}")
            program_lines.append("")
            for theme in (item for item in program["temas"] if item["bloque"] == block):
                program_lines.append(
                    f"{theme['numero_oficial']}. **{theme['id']}** · {theme['titulo']}"
                )
            program_lines.append("")
        markdown = markdown.replace("{{PROGRAMA_MAESTRO}}", "\n".join(program_lines))
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(doc, running)
    add_support_cover(doc, eyebrow, title, "Material de preparación integral y controlado")
    bm.add_markdown(doc, markdown)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return {
        "key": key,
        "status": "generated",
        "source": str(source.relative_to(ROOT)),
        "output": str(output.relative_to(ROOT)),
        "words": len(re.findall(r"\b[\wÁÉÍÓÚÜÑáéíóúüñ]+\b", markdown)),
    }


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def inspect_pdf(path: Path) -> dict:
    reader = PdfReader(str(path), strict=False)
    pages = len(reader.pages)
    if pages:
        _ = reader.pages[0].mediabox
        _ = reader.pages[-1].mediabox
    return {
        "relative_path": str(path.relative_to(ROOT)),
        "category": path.parent.name,
        "filename": path.name,
        "pages": pages,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "validated": pages > 0,
    }


def build_official_archive_index() -> dict:
    pdfs = sorted(
        (
            path
            for path in OFFICIAL_ROOT.rglob("*.pdf")
            if path.name != "Guía completa de la oposición.pdf"
        ),
        key=lambda p: str(p).lower(),
    )
    records = [inspect_pdf(path) for path in pdfs]

    json_path = SOURCE_ROOT / "Inventario técnico de fuentes oficiales.json"
    json_payload = {
        "fecha_corte": "2026-07-30",
        "total_archivos": len(records),
        "total_paginas": sum(item["pages"] for item in records),
        "total_bytes": sum(item["bytes"] for item in records),
        "errores": [item for item in records if not item["validated"]],
        "archivos": records,
    }
    json_path.write_text(
        json.dumps(json_payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    output = SOURCE_ROOT / "Índice y ficha técnica del archivo oficial.docx"
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(doc, "Archivo oficial · Validado 30-07-2026")
    add_support_cover(
        doc,
        "ARCHIVO OFICIAL",
        "Índice y ficha técnica de las fuentes descargadas",
        "Trazabilidad, integridad y localización de los documentos oficiales",
    )

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("1. Resultado de la validación")
    summary_rows = [
        ["Archivos PDF", str(len(records))],
        ["Páginas verificadas", f"{sum(item['pages'] for item in records):,}".replace(",", ".")],
        ["Archivos dañados o vacíos", str(sum(not item["validated"] for item in records))],
        ["Huella de integridad", "SHA-256 individual"],
        ["Fecha de control", UPDATE_DATE_TEXT],
    ]
    add_grid_table(doc, ["Magnitud", "Resultado"], summary_rows, [3000, 6026])

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("2. Distribución por carpetas")
    by_folder = Counter(item["category"] for item in records)
    folder_rows = [
        [
            category,
            str(count),
            str(sum(item["pages"] for item in records if item["category"] == category)),
        ]
        for category, count in sorted(by_folder.items())
    ]
    add_grid_table(doc, ["Carpeta", "PDF", "Páginas"], folder_rows, [5200, 1500, 2326])

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("3. Relación completa")
    rows = [
        [
            item["category"],
            item["filename"],
            str(item["pages"]),
            item["sha256"][:16] + "…",
        ]
        for item in records
    ]
    add_grid_table(
        doc,
        ["Carpeta", "Documento", "Pág.", "SHA-256 abreviado"],
        rows,
        [2100, 4400, 800, 1726],
    )

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("4. Criterio de uso")
    paragraphs = [
        (
            "Los cuestionarios, plantillas, resoluciones y notas se conservan en su formato "
            "oficial para permitir la consulta directa. El material didáctico no sustituye "
            "a estos documentos ni a una eventual publicación posterior del órgano convocante."
        ),
        (
            "La huella SHA-256 permite detectar alteraciones accidentales del archivo local. "
            "La ficha JSON adjunta contiene la ruta, el tamaño, el número de páginas y la huella "
            "completa de cada PDF."
        ),
        (
            "Cuando exista diferencia entre un resumen del material y el texto oficial, prevalece "
            "siempre el texto publicado por el BOE o por el órgano de selección."
        ),
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)

    doc.save(output)
    return {
        "output": str(output.relative_to(ROOT)),
        "json": str(json_path.relative_to(ROOT)),
        "files": len(records),
        "pages": sum(item["pages"] for item in records),
        "errors": sum(not item["validated"] for item in records),
    }


def extract_theme_sources(markdown: str) -> list[dict]:
    links = []
    seen = set()
    for label, url in re.findall(r"\[([^\]]+)\]\((https?://[^)]+)\)", markdown):
        key = url.rstrip(".,;")
        if key in seen:
            continue
        seen.add(key)
        links.append({"label": bm.clean_inline_markdown(label).strip(), "url": key})
    for url in re.findall(r"(?<!\()https?://[^\s)>]+", markdown):
        key = url.rstrip(".,;")
        if key in seen:
            continue
        seen.add(key)
        links.append({"label": "Enlace oficial", "url": key})
    return links


def build_source_cards() -> dict:
    program = json.loads(PROGRAM_PATH.read_text(encoding="utf-8"))
    themes = program["temas"]
    output = SOURCE_ROOT / "Fichas de fuentes oficiales por tema.docx"
    json_path = SOURCE_ROOT / "Fuentes oficiales por tema.json"

    cards = []
    for theme in themes:
        source_path = ROOT / "_trabajo" / "contenidos" / f"{theme['id']}.md"
        if source_path.exists():
            markdown = source_path.read_text(encoding="utf-8")
            links = extract_theme_sources(markdown)
            status = "Completa" if links else "Sin URL extraída"
        else:
            links = []
            status = "Pendiente de redacción"
        cards.append(
            {
                "id": theme["id"],
                "bloque": theme["bloque"],
                "numero": theme["numero_oficial"],
                "titulo": theme["titulo"],
                "status": status,
                "sources": links,
            }
        )

    json_path.write_text(
        json.dumps(
            {"fecha_corte": "2026-07-30", "temas": cards},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(doc, "Fichas de fuentes · Corte 30-07-2026")
    add_support_cover(
        doc,
        "FUENTES POR TEMA",
        "Fichas de fuentes oficiales y control de actualización",
        "Trazabilidad temática de las normas y páginas institucionales utilizadas",
    )

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Criterio de la ficha")
    p = doc.add_paragraph(
        "Cada ficha identifica el título literal del programa, el estado de redacción y "
        "los enlaces oficiales citados en el tema. La fecha de corte general es el "
        "30 de julio de 2026. Los cambios posteriores deben registrarse en el control "
        "de actualización antes de incorporarse al material."
    )

    for card in cards:
        p = doc.add_paragraph(style="Heading 1")
        prefix = "Tema" if card["id"].startswith("G") else "Tema de Seguridad Social"
        p.add_run(f"{card['id']} · {prefix} {card['numero']}")

        p = doc.add_paragraph()
        r = p.add_run(card["titulo"])
        bm.set_run_font(r, size=11, bold=True, color=bm.NAVY)

        add_grid_table(
            doc,
            ["Campo", "Dato"],
            [
                ["Bloque", card["bloque"]],
                ["Estado", card["status"]],
                ["Fecha de corte", UPDATE_DATE_TEXT],
                ["Fuentes enlazadas", str(len(card["sources"]))],
            ],
            [2200, 6826],
        )
        if card["sources"]:
            rows = [
                [str(index), source["label"], source["url"]]
                for index, source in enumerate(card["sources"], 1)
            ]
            add_grid_table(
                doc,
                ["N.º", "Fuente", "URL oficial"],
                rows,
                [700, 3000, 5326],
            )
        else:
            p = doc.add_paragraph(
                "Ficha pendiente: el tema todavía no estaba disponible al generar este control."
            )
            bm.add_paragraph_shading(p, "FFF2CC")

    doc.save(output)
    return {
        "output": str(output.relative_to(ROOT)),
        "json": str(json_path.relative_to(ROOT)),
        "themes": len(cards),
        "complete": sum(bool(card["sources"]) for card in cards),
    }


def docx_word_count(path: Path) -> int:
    try:
        doc = Document(path)
    except Exception:
        return 0
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        text += "\n" + "\n".join(
            cell.text for row in table.rows for cell in row.cells
        )
    return len(re.findall(r"\b[\wÁÉÍÓÚÜÑáéíóúüñ]+\b", text))


def docx_integrity(path: Path) -> tuple[bool, str]:
    try:
        with ZipFile(path) as archive:
            required = {"[Content_Types].xml", "word/document.xml"}
            missing = required.difference(archive.namelist())
            bad = archive.testzip()
        if missing:
            return False, "Faltan: " + ", ".join(sorted(missing))
        if bad:
            return False, f"Entrada ZIP dañada: {bad}"
        return True, "Correcto"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def build_control_document() -> dict:
    program = json.loads(PROGRAM_PATH.read_text(encoding="utf-8"))
    expected_theme_ids = [theme["id"] for theme in program["temas"]]

    theme_docs = sorted(THEME_ROOT.rglob("*.docx"))
    test_docs = sorted(TEST_ROOT.rglob("*.docx"))
    sim_without = sorted((SIM_ROOT / "01. Sin respuestas").glob("*.docx"))
    sim_answers = sorted((SIM_ROOT / "02. Plantillas de respuestas").glob("*.docx"))
    sim_reasoned = sorted((SIM_ROOT / "03. Versiones razonadas profesorado").glob("*.docx"))
    all_official_pdfs = sorted(OFFICIAL_ROOT.rglob("*.pdf"))
    guide_pdfs = [path for path in all_official_pdfs if path.name == "Guía completa de la oposición.pdf"]
    official_pdfs = [path for path in all_official_pdfs if path not in guide_pdfs]
    normative_pdfs = sorted(NORM_ROOT.rglob("*.pdf"))
    normative_docx = sorted(NORM_ROOT.rglob("*.docx"))
    analysis_docx = sorted(ANALYSIS_ROOT.rglob("*.docx"))
    analysis_pdfs = sorted(ANALYSIS_ROOT.rglob("*.pdf"))
    analysis_workbooks = sorted(ANALYSIS_ROOT.rglob("*.xlsx"))

    all_docx = sorted(
        set(
            theme_docs
            + test_docs
            + sim_without
            + sim_answers
            + sim_reasoned
            + list(OFFICIAL_ROOT.rglob("*.docx"))
            + normative_docx
            + list(SOURCE_ROOT.rglob("*.docx"))
            + list(REPORT_ROOT.rglob("*.docx"))
            + analysis_docx
        )
    )
    integrity_rows = []
    invalid = 0
    for path in all_docx:
        ok, message = docx_integrity(path)
        invalid += not ok
        integrity_rows.append(
            [
                str(path.relative_to(ROOT)),
                "Sí" if ok else "No",
                f"{docx_word_count(path):,}".replace(",", "."),
                message,
            ]
        )

    output = REPORT_ROOT / "Tabla final de control de entregables.docx"
    json_path = QA_ROOT / "control_final_entregables.json"
    summary = {
        "fecha_corte": "2026-07-30",
        "programa_temas": len(expected_theme_ids),
        "temas_docx": len(theme_docs),
        "tests_docx": len(test_docs),
        "simulacros_sin_respuestas": len(sim_without),
        "plantillas_simulacros": len(sim_answers),
        "simulacros_razonados": len(sim_reasoned),
        "pdf_oficiales": len(official_pdfs),
        "guia_pdf": len(guide_pdfs),
        "pdf_normativa": len(normative_pdfs),
        "docx_normativa": len(normative_docx),
        "informes_analisis_docx": len(analysis_docx),
        "informes_analisis_pdf": len(analysis_pdfs),
        "libros_analisis_xlsx": len(analysis_workbooks),
        "docx_verificados": len(all_docx),
        "docx_invalidos": invalid,
    }
    json_path.write_text(
        json.dumps(
            {
                "resumen": summary,
                "archivos_docx": [
                    {
                        "path": row[0],
                        "integridad": row[1],
                        "palabras": row[2],
                        "resultado": row[3],
                    }
                    for row in integrity_rows
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(doc, "Control de entregables · 30-07-2026")
    add_support_cover(
        doc,
        "CONTROL FINAL",
        "Tabla de control de entregables",
        "Inventario, integridad técnica y estado de completitud",
    )

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("1. Resumen cuantitativo")
    rows = [
        ["Temas redactados", str(summary["temas_docx"]), "36"],
        ["Test por temas", str(summary["tests_docx"]), "36"],
        ["Simulacros sin respuestas", str(summary["simulacros_sin_respuestas"]), "20"],
        ["Plantillas de simulacro", str(summary["plantillas_simulacros"]), "20"],
        ["Simulacros razonados", str(summary["simulacros_razonados"]), "20"],
        ["PDF oficiales archivados", str(summary["pdf_oficiales"]), "Control documental"],
        ["Guía completa en PDF", str(summary["guia_pdf"]), "1"],
        ["PDF de normativa", str(summary["pdf_normativa"]), "Normas utilizadas"],
        ["Documentos Word de normativa", str(summary["docx_normativa"]), "3"],
        ["Informe de análisis en Word", str(summary["informes_analisis_docx"]), "1"],
        ["Informe de análisis en PDF", str(summary["informes_analisis_pdf"]), "1"],
        ["Base de análisis en Excel", str(summary["libros_analisis_xlsx"]), "1"],
        ["DOCX verificados", str(summary["docx_verificados"]), "Todos los DOCX"],
        ["DOCX inválidos", str(summary["docx_invalidos"]), "0"],
    ]
    add_grid_table(doc, ["Entregable", "Disponible", "Objetivo"], rows, [4300, 1800, 2926])

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("2. Integridad de los documentos Word")
    add_grid_table(
        doc,
        ["Ruta", "ZIP/XML", "Palabras", "Resultado"],
        integrity_rows,
        [4550, 1000, 1000, 2476],
    )

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("3. Criterios del cierre")
    for text in (
        "La integridad ZIP/XML confirma que el archivo DOCX puede abrirse y contiene las partes esenciales del formato Office Open XML.",
        "El recuento de palabras es un control auxiliar; no sustituye la inspección visual de la paginación, los saltos, las tablas ni los pies de página.",
        "La tabla debe regenerarse después de cualquier modificación material para que los recuentos y estados reflejen el contenido final.",
    ):
        doc.add_paragraph(text, style="List Bullet")
    doc.save(output)
    return {
        "output": str(output.relative_to(ROOT)),
        "json": str(json_path.relative_to(ROOT)),
        **summary,
    }


def build_legal_notice() -> dict:
    output = OFFICIAL_ROOT / "Aviso legal y condiciones de uso.docx"
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(doc, "Aviso legal · 30-07-2026")
    add_support_cover(
        doc,
        "AVISO LEGAL",
        "Condiciones de uso y propiedad intelectual",
        "Alcance del material, límites y uso autorizado",
    )
    sections = [
        (
            "1. Titularidad",
            [
                "© 2026 ACADEMIA LORMAN. Todos los derechos reservados.",
                "La redacción, selección, organización, esquemas, preguntas, explicaciones y diseño del material constituyen una obra protegida en los términos del Real Decreto Legislativo 1/1996.",
            ],
        ),
        (
            "2. Usos prohibidos",
            [
                "No se autoriza la reproducción, distribución, comunicación pública, transformación, venta, cesión, publicación o difusión total o parcial sin consentimiento previo y escrito de la titular.",
                "Queda expresamente prohibida la publicación en páginas web, plataformas educativas, redes sociales, grupos de mensajería, aplicaciones, repositorios, nubes compartidas o servicios de intercambio de archivos.",
            ],
        ),
        (
            "3. Fuentes oficiales",
            [
                "Los textos normativos, resoluciones, cuestionarios y demás documentos oficiales conservan la titularidad y el régimen jurídico que les corresponda. Se archivan o enlazan para fines de estudio, trazabilidad y verificación.",
                "Ante cualquier divergencia, prevalece el texto publicado por el Boletín Oficial del Estado o por el órgano de selección.",
            ],
        ),
        (
            "4. Actualización y responsabilidad",
            [
                "La fecha de corte general es el 30 de julio de 2026. Las normas y actuaciones posteriores deben comprobarse antes de usar el material en una convocatoria futura.",
                "El material es una herramienta de preparación y no constituye asesoramiento jurídico ni garantía de superación del proceso selectivo.",
            ],
        ),
        (
            "5. Uso personal",
            [
                "La adquisición o entrega de una copia no transmite los derechos de explotación. Se autoriza únicamente el uso personal de estudio por la persona destinataria.",
                "Las copias de seguridad personales deberán mantenerse en un entorno privado y bajo control de la persona usuaria.",
            ],
        ),
    ]
    for heading, paragraphs in sections:
        p = doc.add_paragraph(style="Heading 1")
        p.add_run(heading)
        for text in paragraphs:
            doc.add_paragraph(text)
    doc.save(output)
    return {"output": str(output.relative_to(ROOT))}


def build_all(selected: set[str] | None = None) -> dict:
    ensure_directories()
    report: dict[str, object] = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "date_cutoff": "2026-07-30",
        "documents": [],
    }
    markdown_keys = set(SUPPORT_SOURCES)
    if selected is None:
        selected = markdown_keys | {"archivo", "fuentes", "control", "legal"}
    for key in sorted(markdown_keys & selected):
        report["documents"].append(build_markdown_document(key))
    if "archivo" in selected:
        report["archive"] = build_official_archive_index()
    if "fuentes" in selected:
        report["sources"] = build_source_cards()
    if "legal" in selected:
        report["legal"] = build_legal_notice()
    if "control" in selected:
        report["control"] = build_control_document()

    report_path = QA_ROOT / "informe_materiales_auxiliares.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--only",
        nargs="*",
        choices=sorted(set(SUPPORT_SOURCES) | {"archivo", "fuentes", "control", "legal"}),
        help="Genera únicamente los elementos indicados.",
    )
    args = parser.parse_args()
    selected = set(args.only) if args.only else None
    report = build_all(selected)
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

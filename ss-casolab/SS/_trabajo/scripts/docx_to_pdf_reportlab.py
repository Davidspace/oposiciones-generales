from __future__ import annotations

import argparse
import html
import io
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image as PILImage
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    LongTable,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
FONT_DIR = Path("C:/Windows/Fonts")
NAVY = colors.HexColor("#17365D")
BLUE = colors.HexColor("#2F5597")
MID_BLUE = colors.HexColor("#5B9BD5")
DARK_GRAY = colors.HexColor("#374151")
MID_GRAY = colors.HexColor("#64748B")
LIGHT_BLUE = colors.HexColor("#EAF3F8")
LIGHT_GRAY = colors.HexColor("#F3F4F6")


def register_fonts() -> None:
    candidates = {
        "Arial": "arial.ttf",
        "Arial-Bold": "arialbd.ttf",
        "Arial-Italic": "ariali.ttf",
        "Arial-BoldItalic": "arialbi.ttf",
    }
    for name, filename in candidates.items():
        path = FONT_DIR / filename
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))


def color_from_hex(value: str | None, fallback=colors.black):
    if not value:
        return fallback
    value = str(value).lstrip("#")
    if re.fullmatch(r"[0-9A-Fa-f]{6}", value):
        return colors.HexColor("#" + value)
    return fallback


def run_markup(run) -> str:
    text = html.escape(run.text or "").replace("\n", "<br/>")
    if not text:
        return ""
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.font.color and run.font.color.rgb:
        text = f'<font color="#{run.font.color.rgb}">{text}</font>'
    return text


def paragraph_markup(paragraph: DocxParagraph) -> str:
    parts = [run_markup(run) for run in paragraph.runs]
    value = "".join(parts)
    if not value:
        value = html.escape(paragraph.text or "")
    value = value.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
    return value.strip()


def extract_images(paragraph: DocxParagraph, document: Document) -> list[bytes]:
    result: list[bytes] = []
    for run in paragraph.runs:
        for blip in run._element.xpath(".//a:blip"):
            relation_id = blip.get(qn("r:embed"))
            if not relation_id:
                continue
            part = document.part.related_parts.get(relation_id)
            if part is not None:
                result.append(part.blob)
    return result


def image_flowable(blob: bytes, max_width: float, max_height: float = 170 * mm):
    with PILImage.open(io.BytesIO(blob)) as image:
        width_px, height_px = image.size
    ratio = min(max_width / width_px, max_height / height_px)
    ratio = min(ratio, 1.0)
    flow = Image(io.BytesIO(blob), width=width_px * ratio, height=height_px * ratio)
    flow.hAlign = "CENTER"
    return flow


def make_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "normal": ParagraphStyle(
            "LormanNormal",
            parent=base["BodyText"],
            fontName="Arial",
            fontSize=10.6,
            leading=13.1,
            textColor=DARK_GRAY,
            alignment=TA_JUSTIFY,
            spaceAfter=4.5,
            allowWidows=0,
            allowOrphans=0,
        ),
        "h1": ParagraphStyle(
            "LormanH1",
            parent=base["Heading1"],
            fontName="Arial-Bold",
            fontSize=15,
            leading=18,
            textColor=BLUE,
            spaceBefore=13,
            spaceAfter=7,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "LormanH2",
            parent=base["Heading2"],
            fontName="Arial-Bold",
            fontSize=12.3,
            leading=15,
            textColor=MID_BLUE,
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "LormanH3",
            parent=base["Heading3"],
            fontName="Arial-Bold",
            fontSize=11.3,
            leading=14,
            textColor=BLUE,
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "h4": ParagraphStyle(
            "LormanH4",
            parent=base["Heading4"],
            fontName="Arial-Bold",
            fontSize=10.7,
            leading=13,
            textColor=DARK_GRAY,
            spaceBefore=6,
            spaceAfter=3,
            keepWithNext=True,
        ),
        "question": ParagraphStyle(
            "LormanQuestion",
            parent=base["BodyText"],
            fontName="Arial-Bold",
            fontSize=10.6,
            leading=13,
            textColor=DARK_GRAY,
            spaceBefore=8,
            spaceAfter=3,
            keepWithNext=True,
        ),
        "bullet": ParagraphStyle(
            "LormanBullet",
            parent=base["BodyText"],
            fontName="Arial",
            fontSize=10.5,
            leading=13,
            textColor=DARK_GRAY,
            leftIndent=8 * mm,
            firstLineIndent=-4 * mm,
            spaceAfter=3,
        ),
        "cover_big": ParagraphStyle(
            "LormanCoverBig",
            parent=base["Title"],
            fontName="Arial-Bold",
            fontSize=20,
            leading=24,
            textColor=NAVY,
            spaceBefore=4,
            spaceAfter=12,
            keepWithNext=True,
        ),
        "cover_mid": ParagraphStyle(
            "LormanCoverMid",
            parent=base["BodyText"],
            fontName="Arial-Bold",
            fontSize=13,
            leading=16,
            textColor=BLUE,
            spaceAfter=6,
        ),
        "cover_small": ParagraphStyle(
            "LormanCoverSmall",
            parent=base["BodyText"],
            fontName="Arial-Bold",
            fontSize=9,
            leading=11,
            textColor=MID_BLUE,
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "LormanSmall",
            parent=base["BodyText"],
            fontName="Arial",
            fontSize=8.3,
            leading=10,
            textColor=MID_GRAY,
            alignment=TA_JUSTIFY,
            spaceAfter=3,
        ),
        "caption": ParagraphStyle(
            "LormanCaption",
            parent=base["BodyText"],
            fontName="Arial-Italic",
            fontSize=8.5,
            leading=10,
            textColor=MID_GRAY,
            alignment=TA_CENTER,
            spaceAfter=7,
        ),
    }


def paragraph_style_for(
    paragraph: DocxParagraph,
    styles: dict[str, ParagraphStyle],
    paragraph_index: int,
) -> ParagraphStyle:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name == "Heading 1":
        return styles["h1"]
    if style_name == "Heading 2":
        return styles["h2"]
    if style_name == "Heading 3":
        return styles["h3"]
    if style_name == "Heading 4":
        return styles["h4"]
    if style_name == "Question":
        return styles["question"]
    if style_name in {"List Bullet", "List Number", "List Paragraph"}:
        return styles["bullet"]
    if paragraph_index < 18 and paragraph.runs:
        max_size = max(
            (
                run.font.size.pt
                for run in paragraph.runs
                if run.font.size is not None
            ),
            default=0,
        )
        if max_size >= 18:
            return styles["cover_big"]
        if max_size >= 12:
            return styles["cover_mid"]
        if all(run.bold for run in paragraph.runs if run.text.strip()):
            return styles["cover_small"]
    if paragraph.runs:
        explicit_sizes = [
            run.font.size.pt for run in paragraph.runs if run.font.size is not None
        ]
        if explicit_sizes and max(explicit_sizes) <= 8.6:
            return styles["small"]
        if all(run.italic for run in paragraph.runs if run.text.strip()):
            return styles["caption"]
    return styles["normal"]


def contains_page_break(paragraph: DocxParagraph) -> bool:
    return bool(paragraph._element.xpath('.//w:br[@w:type="page"]'))


def convert_table(
    table: DocxTable,
    available_width: float,
    styles: dict[str, ParagraphStyle],
):
    rows: list[list[Paragraph]] = []
    max_cols = max((len(row.cells) for row in table.rows), default=1)
    for row in table.rows:
        values: list[Paragraph] = []
        for cell in row.cells:
            text = "\n".join(
                re.sub(r"\s+", " ", paragraph.text).strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            values.append(Paragraph(html.escape(text).replace("\n", "<br/>"), styles["small"]))
        while len(values) < max_cols:
            values.append(Paragraph("", styles["small"]))
        rows.append(values)
    if not rows:
        return Spacer(1, 1)
    col_widths = [available_width / max_cols] * max_cols
    result = LongTable(
        rows,
        colWidths=col_widths,
        repeatRows=1,
        splitByRow=1,
        hAlign="LEFT",
    )
    result.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_BLUE),
                ("TEXTCOLOR", (0, 0), (-1, 0), NAVY),
                ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Arial"),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#B8C7D9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
            ]
        )
    )
    return result


def extract_running_label(document: Document) -> str:
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = re.sub(r"\s+", " ", paragraph.text).strip()
            if text:
                return text
    return "Material de preparación · Actualizado 30-07-2026"


def make_page_drawer(running_label: str):
    legal = (
        "© 2026 ACADEMIA LORMAN. Todos los derechos reservados. "
        "Prohibida su reproducción, distribución o difusión sin autorización escrita."
    )

    def draw(canvas, doc):
        canvas.saveState()
        page = canvas.getPageNumber()
        width, height = A4
        if page > 1:
            canvas.setStrokeColor(colors.HexColor("#B4C7E7"))
            canvas.setLineWidth(0.6)
            canvas.line(25.4 * mm, height - 17 * mm, width - 25.4 * mm, height - 17 * mm)
            canvas.setFont("Arial-Bold", 7.4)
            canvas.setFillColor(BLUE)
            canvas.drawString(25.4 * mm, height - 14 * mm, running_label[:115])
        canvas.setFont("Arial", 6.2)
        canvas.setFillColor(MID_GRAY)
        canvas.drawString(25.4 * mm, 12.5 * mm, legal)
        canvas.setFont("Arial", 7)
        canvas.drawRightString(width - 25.4 * mm, 9.2 * mm, f"Página {page}")
        canvas.restoreState()

    return draw


def convert(input_path: Path, output_path: Path) -> dict:
    register_fonts()
    document = Document(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = make_styles()
    page_width, page_height = A4
    left = right = 25.4 * mm
    top = 22 * mm
    bottom = 22 * mm
    available_width = page_width - left - right

    pdf = BaseDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=left,
        rightMargin=right,
        topMargin=top,
        bottomMargin=bottom,
        title=document.core_properties.title or input_path.stem,
        author="ACADEMIA LORMAN",
        subject=document.core_properties.subject or "",
    )
    frame = Frame(
        left,
        bottom,
        available_width,
        page_height - top - bottom,
        id="body",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    running_label = extract_running_label(document)
    template = PageTemplate(
        id="lorman",
        frames=[frame],
        onPage=make_page_drawer(running_label),
    )
    pdf.addPageTemplates([template])

    story = []
    paragraph_index = 0
    list_counter = 0
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, document)
            style_name = paragraph.style.name if paragraph.style else ""
            markup = paragraph_markup(paragraph)
            images = extract_images(paragraph, document)

            if contains_page_break(paragraph):
                if markup:
                    style = paragraph_style_for(paragraph, styles, paragraph_index)
                    story.append(Paragraph(markup, style))
                story.append(PageBreak())
                paragraph_index += 1
                list_counter = 0
                continue

            if style_name == "List Number":
                list_counter += 1
                markup = f"<b>{list_counter}.</b>&nbsp; {markup}"
            elif style_name in {"List Bullet", "List Paragraph"}:
                list_counter = 0
                markup = f"•&nbsp; {markup}"
            else:
                list_counter = 0

            if markup:
                style = paragraph_style_for(paragraph, styles, paragraph_index)
                flow = Paragraph(markup, style)
                if style_name == "Question":
                    story.append(KeepTogether([flow]))
                else:
                    story.append(flow)
            for blob in images:
                story.append(Spacer(1, 3))
                story.append(image_flowable(blob, available_width))
                story.append(Spacer(1, 5))
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            table = DocxTable(child, document)
            story.append(convert_table(table, available_width, styles))
            story.append(Spacer(1, 6))
            list_counter = 0

    pdf.build(story)
    reader = PdfReader(str(output_path))
    return {
        "input": str(input_path.relative_to(ROOT)),
        "output": str(output_path.relative_to(ROOT)),
        "pages": len(reader.pages),
        "bytes": output_path.stat().st_size,
        "encrypted": reader.is_encrypted,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    input_path = args.input if args.input.is_absolute() else ROOT / args.input
    output_path = args.output if args.output.is_absolute() else ROOT / args.output
    print(json.dumps(convert(input_path, output_path), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

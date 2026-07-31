from __future__ import annotations

import json
import random
import re
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

import build_materials as bm  # noqa: E402
import build_supporting_materials as bs  # noqa: E402


ROOT = SCRIPT_DIR.parents[1]
PROGRAM_PATH = ROOT / "_trabajo" / "investigacion" / "programa_maestro.json"
TEST_DIR = ROOT / "_trabajo" / "tests"
SIM_ROOT = ROOT / "04. Simulacros"
QA_DIR = ROOT / "_trabajo" / "qa"
UPDATE_DATE = "30 de julio de 2026"
UPDATE_NUMERIC = "30-07-2026"

STUDENT_DIR = SIM_ROOT / "01. Sin respuestas"
ANSWER_DIR = SIM_ROOT / "02. Plantillas de respuestas"
TEACHER_DIR = SIM_ROOT / "03. Versiones razonadas profesorado"


@dataclass(frozen=True)
class Question:
    source_id: str
    theme_id: str
    theme_number: int
    block: str
    source_number: int
    stem: str
    options: tuple[str, str, str, str]
    correct: int
    difficulty: str
    verification: str


THEORY_PREFIXES = [
    "Conforme al ordenamiento vigente a la fecha de corte,",
    "A efectos de aplicar correctamente la regla examinada,",
    "En una comprobación jurídica del expediente,",
    "Sin añadir circunstancias distintas de las indicadas,",
    "De acuerdo con la regulación aplicable,",
    "Para resolver correctamente la cuestión planteada,",
    "En relación con el régimen jurídico vigente,",
    "Durante el control de legalidad de una actuación administrativa,",
    "Al verificar el cumplimiento de la norma,",
    "En el marco del programa oficial de acceso libre,",
    "Con arreglo a la redacción consolidada aplicable,",
    "Ante una consulta formulada a una unidad administrativa,",
    "Para determinar la consecuencia jurídica correcta,",
    "En una revisión de conceptos próximos que no deben confundirse,",
    "Atendiendo a la regla general y a sus excepciones expresas,",
    "En una actuación sometida al Derecho público,",
    "Para identificar el órgano, plazo o efecto jurídicamente correcto,",
    "En la resolución de una incidencia de tramitación,",
    "A la vista de la normativa oficial utilizada en el temario,",
    "En una prueba de precisión normativa,",
    "Al contrastar las cuatro alternativas con la norma,",
    "En un supuesto de aplicación ordinaria de la regulación vigente,",
    "Para evitar una atribución competencial incorrecta,",
    "En una consulta sobre el alcance exacto del precepto,",
]

PRACTICAL_PREFIXES = [
    "En el subexpediente descrito y con los únicos datos facilitados,",
    "La unidad competente debe resolver la siguiente cuestión:",
    "Al comprobar el encuadramiento y sus efectos,",
    "En la tramitación del expediente de Seguridad Social,",
    "La persona funcionaria responsable necesita determinar si",
    "Antes de dictar resolución, debe aplicarse la regla según la cual",
    "En la revisión del expediente económico-administrativo,",
    "Para informar correctamente a la persona interesada,",
    "Al calcular plazos, requisitos o efectos,",
    "En la gestión coordinada entre INSS y TGSS,",
    "Sin presumir hechos no incluidos en el caso,",
    "En la fase procedimental correspondiente,",
    "Al verificar el derecho o la obligación de Seguridad Social,",
    "La dirección provincial plantea la siguiente comprobación:",
    "En el control de la prestación o de la cotización,",
    "Para decidir qué entidad u órgano resulta competente,",
    "A efectos de reconocer, suspender o extinguir el derecho,",
    "En el expediente individualizado de la persona afectada,",
    "Al aplicar la norma especial frente a la regla general,",
    "En la propuesta de resolución que se está preparando,",
    "Para resolver una discrepancia entre los datos declarados y los acreditados,",
    "En el trámite de audiencia y comprobación del expediente,",
    "Al analizar la responsabilidad de ingreso o de prestación,",
    "En una consulta integrada de afiliación, cotización y prestaciones,",
]


def clean_md(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    return re.sub(r"\s+", " ", text).strip()


def parse_test(path: Path, theme: dict) -> list[Question]:
    lines = path.read_text(encoding="utf-8").splitlines()
    questions: list[Question] = []
    current_number: int | None = None
    current_stem = ""
    options: list[str] = []
    correct = -1

    for raw in lines:
        stripped = raw.strip()
        # Los bancos temáticos admiten la pregunta como párrafo numerado
        # o como encabezado Markdown (p. ej., ``### 1. Enunciado``).
        question_match = re.match(r"^(?:#{1,6}\s*)?(\d+)\.\s+(.+)$", stripped)
        if question_match:
            current_number = int(question_match.group(1))
            current_stem = clean_md(question_match.group(2))
            options = []
            correct = -1
            continue

        option_match = re.match(
            r"^-\s+(\*\*)?([A-Da-d])\)\s+(.+?)(\*\*)?$", stripped
        )
        if current_number is not None and option_match:
            is_correct = bool(option_match.group(1) and option_match.group(4))
            option_text = clean_md(option_match.group(3))
            if is_correct:
                correct = len(options)
            options.append(option_text)
            continue

        comment_match = re.search(
            r"verificaci[oó]n:\s*(.*?);\s*dificultad:\s*(b[aá]sica|media|alta)",
            stripped,
            flags=re.IGNORECASE,
        )
        if current_number is not None and comment_match:
            difficulty = (
                comment_match.group(2).lower().replace("basica", "básica")
            )
            if len(options) != 4 or correct not in range(4):
                raise ValueError(f"Pregunta inválida en {path.name}: {current_number}")
            questions.append(
                Question(
                    source_id=f"{theme['id']}-{current_number:02d}",
                    theme_id=theme["id"],
                    theme_number=theme["numero_oficial"],
                    block=theme["bloque"],
                    source_number=current_number,
                    stem=current_stem,
                    options=tuple(options),  # type: ignore[arg-type]
                    correct=correct,
                    difficulty=difficulty,
                    verification=clean_md(comment_match.group(1)),
                )
            )
            current_number = None
            current_stem = ""
            options = []
            correct = -1

    if len(questions) != 30:
        raise ValueError(f"{path.name}: se esperaban 30 preguntas y hay {len(questions)}")
    return questions


def load_bank() -> tuple[dict[str, dict], dict[str, list[Question]]]:
    program = json.loads(PROGRAM_PATH.read_text(encoding="utf-8"))
    themes = {item["id"]: item for item in program["temas"]}
    bank: dict[str, list[Question]] = {}
    missing = []
    for theme_id, theme in themes.items():
        path = TEST_DIR / f"{theme_id}.md"
        if not path.exists():
            missing.append(theme_id)
            continue
        bank[theme_id] = parse_test(path, theme)
    if missing:
        raise RuntimeError(
            "Faltan test temáticos para construir los simulacros: " + ", ".join(missing)
        )
    return themes, bank


def difficulty_plan(level: str, part: str) -> list[str]:
    if part == "theory":
        distributions = {
            "básico-medio": {"básica": 30, "media": 34, "alta": 9},
            "nivel oficial": {"básica": 20, "media": 36, "alta": 17},
            "alta": {"básica": 10, "media": 30, "alta": 33},
        }
    else:
        distributions = {
            "básico-medio": {"básica": 7, "media": 9, "alta": 2},
            "nivel oficial": {"básica": 4, "media": 9, "alta": 5},
            "alta": {"básica": 2, "media": 7, "alta": 9},
        }
    return [
        difficulty
        for difficulty, count in distributions[level].items()
        for _ in range(count)
    ]


def answer_plan(part: str) -> list[int]:
    counts = [19, 18, 18, 18] if part == "theory" else [5, 5, 4, 4]
    return [letter for letter, count in enumerate(counts) for _ in range(count)]


def theme_quota(ids: list[str], total: int, rng: random.Random) -> list[str]:
    if total < len(ids):
        raise ValueError("El total debe permitir al menos una pregunta por tema.")
    result = list(ids)
    while len(result) < total:
        result.append(ids[rng.randrange(len(ids))])
    rng.shuffle(result)
    return result


def pick_questions(
    theme_ids: list[str],
    desired_difficulties: list[str],
    bank: dict[str, list[Question]],
    rng: random.Random,
    excluded_source_ids: set[str],
) -> list[Question]:
    if len(theme_ids) != len(desired_difficulties):
        raise ValueError("Cuotas y dificultades no coinciden.")
    picked: list[Question] = []
    used = set(excluded_source_ids)
    for theme_id, difficulty in zip(theme_ids, desired_difficulties):
        candidates = [
            q
            for q in bank[theme_id]
            if q.difficulty == difficulty and q.source_id not in used
        ]
        if not candidates:
            candidates = [q for q in bank[theme_id] if q.source_id not in used]
        if not candidates:
            raise RuntimeError(f"No quedan preguntas disponibles para {theme_id}")
        selected = candidates[rng.randrange(len(candidates))]
        picked.append(selected)
        used.add(selected.source_id)
    return picked


def rotate_options(question: Question, desired_correct: int) -> tuple[list[str], int]:
    correct_text = question.options[question.correct]
    distractors = [
        value for idx, value in enumerate(question.options) if idx != question.correct
    ]
    result: list[str] = []
    distractor_index = 0
    for idx in range(4):
        if idx == desired_correct:
            result.append(correct_text)
        else:
            result.append(distractors[distractor_index])
            distractor_index += 1
    return result, desired_correct


def lower_after_prefix(stem: str) -> str:
    if stem.startswith("¿") or not stem:
        return stem
    return stem[0].lower() + stem[1:]


def variant_stem(
    question: Question,
    *,
    practical: bool,
    variant_index: int,
) -> str:
    # Dos preguntas de origen abordaban el mismo concepto con un enunciado
    # literalmente idéntico desde perspectivas temáticas distintas. Se
    # contextualiza una de ellas para que la batería completa no contenga
    # duplicados exactos.
    source_stem = {
        "G21-10": (
            "En el marco del régimen jurídico de la discapacidad, "
            "¿qué se entiende por ajuste razonable?"
        )
    }.get(question.source_id, question.stem)
    prefixes = PRACTICAL_PREFIXES if practical else THEORY_PREFIXES
    prefix = prefixes[variant_index % len(prefixes)]
    stem = lower_after_prefix(source_stem)
    if prefix.endswith(":"):
        if stem.startswith("¿"):
            return f"{prefix} {stem}"
        return f"{prefix} {source_stem}"
    return f"{prefix} {stem}"


def level_for_sim(number: int) -> str:
    if number <= 5:
        return "básico-medio"
    if number <= 15:
        return "nivel oficial"
    return "alta"


def add_exam_cover(doc: Document, number: int, level: str, kind: str) -> None:
    title = (
        f"Simulacro {number:02d} · {level}"
        if kind == "student"
        else (
            f"Plantilla de respuestas · Simulacro {number:02d}"
            if kind == "answers"
            else f"Versión razonada del profesorado · Simulacro {number:02d}"
        )
    )
    bm.add_cover(
        doc,
        kind="Simulacro completo",
        number=None,
        title=title,
        block="Cuerpo Administrativo de la Administración de la Seguridad Social · C1 · acceso libre",
    )


def add_instructions(doc: Document, number: int, level: str, kind: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Instrucciones y estructura oficial reproducida")
    rows = [
        ["Sistema selectivo", "Oposición · ejercicio único con dos partes obligatorias y eliminatorias"],
        ["Tiempo conjunto", "120 minutos"],
        ["Primera parte", "70 preguntas + 3 de reserva · anexos I.A y I.B · 0 a 50 puntos"],
        ["Segunda parte", "Supuesto práctico de 15 preguntas + 3 de reserva · anexo I.B · 0 a 50 puntos"],
        ["Respuesta incorrecta", "Penaliza un cuarto del valor de una respuesta correcta"],
        ["Sin contestar", "No resta"],
        ["Mínimo", "25 puntos transformados en cada parte; la segunda solo se corrige tras superar la primera"],
        ["Nivel de este simulacro", level],
        ["Fecha de actualización", UPDATE_DATE],
    ]
    bs.add_grid_table(doc, ["Elemento", "Regla"], rows, [2600, 6426])
    if kind == "student":
        p = doc.add_paragraph()
        bm.add_paragraph_shading(p, "FFF2CC")
        bm.add_inline_runs(
            p,
            "**No consulte la plantilla durante la realización.** Controle el tiempo de forma conjunta y reserve unos minutos para revisar omisiones y traslados de respuesta.",
            default_size=10,
        )
    elif kind == "teacher":
        p = doc.add_paragraph()
        bm.add_paragraph_shading(p, bm.LIGHT_BLUE)
        bm.add_inline_runs(
            p,
            "Cada respuesta incorpora el tema de procedencia y la norma o precepto que la justifica. La explicación se refiere a la fecha de corte 30-07-2026.",
            default_size=10,
        )


def add_question(
    doc: Document,
    number: int,
    stem: str,
    options: list[str],
    correct: int,
    *,
    show_answer: bool,
    explanation: str | None = None,
) -> None:
    p = doc.add_paragraph(style="Question")
    p.paragraph_format.keep_with_next = True
    bm.add_inline_runs(p, f"{number}. {stem}")
    for idx, option in enumerate(options):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{'ABCD'[idx]}) {option}")
        bm.set_run_font(run, size=10.5, bold=(show_answer and idx == correct))
    if show_answer and explanation:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.keep_together = True
        bm.add_paragraph_shading(p, bm.LIGHT_GRAY)
        lead = p.add_run(f"Respuesta {'ABCD'[correct]} · ")
        bm.set_run_font(lead, size=8.5, bold=True, color=bm.BLUE)
        body = p.add_run(explanation)
        bm.set_run_font(body, size=8.5, color=bm.DARK_GRAY)


def add_response_grid(
    doc: Document,
    title: str,
    count: int,
    *,
    answers: list[int] | None = None,
) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(title)
    rows = []
    columns_per_row = 4
    for start in range(1, count + 1, columns_per_row):
        row: list[str] = []
        for number in range(start, min(start + columns_per_row, count + 1)):
            if answers is None:
                value = f"{number:02d}:  ☐ A   ☐ B   ☐ C   ☐ D"
            else:
                value = f"{number:02d}:  {'ABCD'[answers[number - 1]]}"
            row.append(value)
        while len(row) < columns_per_row:
            row.append("")
        rows.append(row)
    bs.add_grid_table(
        doc,
        ["", "", "", ""],
        rows,
        [2256, 2256, 2256, 2258],
    )


def practical_intro(number: int) -> str:
    names = [
        "Dirección Provincial integrada",
        "Administración de la Tesorería General",
        "Centro de atención e información del INSS",
        "unidad conjunta de encuadramiento y prestaciones",
        "equipo de revisión de expedientes",
    ]
    contexts = [
        "atiende a varias personas trabajadoras, pensionistas y empresas con incidencias conectadas",
        "revisa inscripciones, altas, liquidaciones, deudas y posibles prestaciones",
        "tramita una serie de subexpedientes independientes que deben resolverse con la normativa vigente",
        "contrasta datos declarados con la información obrante en el sistema",
        "debe informar y proponer resolución sin atribuir al INSS funciones de la TGSS ni a la inversa",
    ]
    return (
        f"La {names[(number - 1) % len(names)]} {contexts[(number * 2) % len(contexts)]}. "
        "Salvo indicación expresa en una pregunta, cada subexpediente es independiente y solo deben utilizarse los datos facilitados. "
        "Resuelva las preguntas aplicando el programa específico de Seguridad Social y la normativa vigente a 30 de julio de 2026."
    )


def make_simulation(
    number: int,
    themes: dict[str, dict],
    bank: dict[str, list[Question]],
    variant_counts: defaultdict[str, int],
) -> tuple[dict, list[dict], list[dict]]:
    level = level_for_sim(number)
    rng = random.Random(20260729 + number * 7919)
    all_ids = list(themes)
    ss_ids = [theme_id for theme_id in all_ids if theme_id.startswith("S")]

    theory_theme_ids = theme_quota(all_ids, 73, rng)
    theory_diffs = difficulty_plan(level, "theory")
    rng.shuffle(theory_diffs)
    theory = pick_questions(
        theory_theme_ids, theory_diffs, bank, rng, excluded_source_ids=set()
    )

    practical_theme_ids = theme_quota(ss_ids, 18, rng)
    practical_diffs = difficulty_plan(level, "practical")
    rng.shuffle(practical_diffs)
    practical = pick_questions(
        practical_theme_ids,
        practical_diffs,
        bank,
        rng,
        excluded_source_ids={q.source_id for q in theory},
    )

    theory_answers = answer_plan("theory")
    practical_answers = answer_plan("practical")
    rng.shuffle(theory_answers)
    rng.shuffle(practical_answers)

    rendered_theory = []
    for question, desired in zip(theory, theory_answers):
        variant = variant_counts[question.source_id]
        variant_counts[question.source_id] += 1
        options, correct = rotate_options(question, desired)
        rendered_theory.append(
            {
                "question": question,
                "stem": variant_stem(
                    question, practical=False, variant_index=variant
                ),
                "options": options,
                "correct": correct,
            }
        )

    rendered_practical = []
    for question, desired in zip(practical, practical_answers):
        variant = variant_counts[question.source_id]
        variant_counts[question.source_id] += 1
        options, correct = rotate_options(question, desired)
        rendered_practical.append(
            {
                "question": question,
                "stem": variant_stem(
                    question, practical=True, variant_index=variant
                ),
                "options": options,
                "correct": correct,
            }
        )

    summary = {
        "number": number,
        "level": level,
        "theory": {
            "questions": 73,
            "ordinary": 70,
            "reserve": 3,
            "difficulty": dict(Counter(q["question"].difficulty for q in rendered_theory)),
            "answers": dict(Counter("ABCD"[q["correct"]] for q in rendered_theory)),
            "themes": dict(Counter(q["question"].theme_id for q in rendered_theory)),
        },
        "practical": {
            "questions": 18,
            "ordinary": 15,
            "reserve": 3,
            "difficulty": dict(Counter(q["question"].difficulty for q in rendered_practical)),
            "answers": dict(Counter("ABCD"[q["correct"]] for q in rendered_practical)),
            "themes": dict(Counter(q["question"].theme_id for q in rendered_practical)),
        },
    }
    return summary, rendered_theory, rendered_practical


def build_student(
    number: int,
    level: str,
    theory: list[dict],
    practical: list[dict],
) -> Path:
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(
        doc, f"Simulacro {number:02d} · Sin respuestas · {UPDATE_NUMERIC}"
    )
    add_exam_cover(doc, number, level, "student")
    add_instructions(doc, number, level, "student")

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Primera parte · cuestionario teórico")
    for idx, item in enumerate(theory, 1):
        if idx == 71:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run("Preguntas de reserva de la primera parte")
        add_question(
            doc,
            idx,
            item["stem"],
            item["options"],
            item["correct"],
            show_answer=False,
        )

    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Segunda parte · supuesto práctico")
    p = doc.add_paragraph(practical_intro(number))
    bm.add_paragraph_shading(p, bm.LIGHT_BLUE)
    for idx, item in enumerate(practical, 1):
        if idx == 16:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run("Preguntas de reserva de la segunda parte")
        add_question(
            doc,
            idx,
            item["stem"],
            item["options"],
            item["correct"],
            show_answer=False,
        )

    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Hoja final de respuestas")
    add_response_grid(doc, "Primera parte · 1 a 73", 73)
    add_response_grid(doc, "Segunda parte · 1 a 18", 18)
    p = doc.add_paragraph()
    bm.add_inline_runs(
        p,
        "Puntuación directa orientativa por parte: **aciertos − errores/4**. Las reservas solo sustituyen preguntas anuladas por orden correlativo. La puntuación oficial es la transformada conforme al baremo del tribunal.",
        default_size=9.5,
    )

    output = STUDENT_DIR / f"Simulacro {number:02d} - {level} - sin respuestas.docx"
    doc.save(output)
    return output


def build_answers(
    number: int,
    level: str,
    theory: list[dict],
    practical: list[dict],
) -> Path:
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(
        doc, f"Plantilla · Simulacro {number:02d} · {UPDATE_NUMERIC}"
    )
    add_exam_cover(doc, number, level, "answers")
    add_instructions(doc, number, level, "answers")
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Plantilla correctora")
    add_response_grid(
        doc,
        "Primera parte · 70 preguntas + 3 reservas",
        73,
        answers=[item["correct"] for item in theory],
    )
    add_response_grid(
        doc,
        "Segunda parte · 15 preguntas + 3 reservas",
        18,
        answers=[item["correct"] for item in practical],
    )
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Corrección")
    rows = [
        ["Aciertos", "Número de respuestas correctas computables"],
        ["Errores", "Número de respuestas incorrectas computables"],
        ["Blancos", "No restan"],
        ["Puntuación directa", "Aciertos − (errores ÷ 4)"],
        ["Reservas", "Solo sustituyen anuladas en el orden indicado por el tribunal"],
        ["Puntuación oficial", "Transformada a 0–50 en cada parte según baremo del tribunal"],
    ]
    bs.add_grid_table(doc, ["Concepto", "Aplicación"], rows, [2800, 6226])
    output = ANSWER_DIR / f"Plantilla Simulacro {number:02d} - {level}.docx"
    doc.save(output)
    return output


def build_teacher(
    number: int,
    level: str,
    theory: list[dict],
    practical: list[dict],
) -> Path:
    doc = Document()
    bm.configure_document(doc)
    bm.configure_headers_and_footers(
        doc, f"Profesorado · Simulacro {number:02d} · {UPDATE_NUMERIC}"
    )
    add_exam_cover(doc, number, level, "teacher")
    add_instructions(doc, number, level, "teacher")

    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Primera parte · solución razonada")
    for idx, item in enumerate(theory, 1):
        if idx == 71:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run("Reservas de la primera parte")
        question: Question = item["question"]
        explanation = (
            f"Tema {question.theme_id} · dificultad {question.difficulty}. "
            f"La opción indicada coincide con {question.verification}. "
            "Las demás alternativas alteran el órgano, el plazo, el requisito, el efecto o la excepción aplicable."
        )
        add_question(
            doc,
            idx,
            item["stem"],
            item["options"],
            item["correct"],
            show_answer=True,
            explanation=explanation,
        )

    doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Segunda parte · solución razonada del supuesto")
    p = doc.add_paragraph(practical_intro(number))
    bm.add_paragraph_shading(p, bm.LIGHT_BLUE)
    for idx, item in enumerate(practical, 1):
        if idx == 16:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run("Reservas de la segunda parte")
        question = item["question"]
        explanation = (
            f"Tema {question.theme_id} · dificultad {question.difficulty}. "
            f"Fundamento: {question.verification}. "
            "La solución respeta la distribución competencial y las condiciones del supuesto; no se presumen datos no facilitados."
        )
        add_question(
            doc,
            idx,
            item["stem"],
            item["options"],
            item["correct"],
            show_answer=True,
            explanation=explanation,
        )

    output = TEACHER_DIR / f"Simulacro {number:02d} - {level} - versión razonada profesorado.docx"
    doc.save(output)
    return output


def source_exact_matches(
    rendered: list[dict],
    bank: dict[str, list[Question]],
) -> int:
    originals = {q.stem for questions in bank.values() for q in questions}
    return sum(item["stem"] in originals for item in rendered)


def build_all() -> dict:
    for path in (STUDENT_DIR, ANSWER_DIR, TEACHER_DIR, QA_DIR):
        path.mkdir(parents=True, exist_ok=True)
    themes, bank = load_bank()
    variant_counts: defaultdict[str, int] = defaultdict(int)
    report = {
        "fecha_corte": "2026-07-30",
        "official_format": {
            "time_minutes": 120,
            "part1": {"ordinary": 70, "reserve": 3},
            "part2": {"ordinary": 15, "reserve": 3},
            "wrong_penalty": "-1/4",
            "blank_penalty": 0,
        },
        "simulations": [],
        "files": [],
    }
    all_stems: list[str] = []

    for number in range(1, 21):
        summary, theory, practical = make_simulation(
            number, themes, bank, variant_counts
        )
        student = build_student(number, summary["level"], theory, practical)
        answers = build_answers(number, summary["level"], theory, practical)
        teacher = build_teacher(number, summary["level"], theory, practical)
        rendered = theory + practical
        summary["source_exact_matches"] = source_exact_matches(rendered, bank)
        summary["files"] = [
            str(student.relative_to(ROOT)),
            str(answers.relative_to(ROOT)),
            str(teacher.relative_to(ROOT)),
        ]
        report["simulations"].append(summary)
        report["files"].extend(summary["files"])
        all_stems.extend(item["stem"] for item in rendered)

    report["audit"] = {
        "simulations": 20,
        "student_versions": len(list(STUDENT_DIR.glob("*.docx"))),
        "answer_templates": len(list(ANSWER_DIR.glob("*.docx"))),
        "teacher_versions": len(list(TEACHER_DIR.glob("*.docx"))),
        "questions_total_including_reserves": 20 * 91,
        "unique_rendered_stems": len(set(all_stems)),
        "duplicate_rendered_stems": len(all_stems) - len(set(all_stems)),
        "exact_matches_with_topic_test_stems": sum(
            sim["source_exact_matches"] for sim in report["simulations"]
        ),
        "source_question_usage_max": max(variant_counts.values()),
        "source_question_usage_distribution": dict(
            Counter(variant_counts.values())
        ),
    }
    invalid = []
    for relative in report["files"]:
        path = ROOT / relative
        ok, message = bs.docx_integrity(path)
        if not ok:
            invalid.append({"path": relative, "message": message})
    report["audit"]["invalid_docx"] = invalid
    report_path = QA_DIR / "auditoria_simulacros.json"
    report_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def main() -> None:
    report = build_all()
    print(
        json.dumps(
            {
                "simulations": len(report["simulations"]),
                **report["audit"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()

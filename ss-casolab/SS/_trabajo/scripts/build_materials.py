from __future__ import annotations

import argparse
import json
import re
import unicodedata
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PROGRAM_PATH = ROOT / "_trabajo" / "investigacion" / "programa_maestro.json"
CONTENT_DIR = ROOT / "_trabajo" / "contenidos"
TEST_DIR = ROOT / "_trabajo" / "tests"
QA_DIR = ROOT / "_trabajo" / "qa"

UPDATE_DATE_TEXT = "30 de julio de 2026"
UPDATE_DATE_NUMERIC = "30-07-2026"
AUTHOR = "ACADEMIA LORMAN"
COPYRIGHT_LINE = f"© 2026 {AUTHOR}. Todos los derechos reservados."
LEGAL_PARAGRAPHS = [
    COPYRIGHT_LINE,
    (
        "Material protegido por el Real Decreto Legislativo 1/1996, de 12 de abril, "
        "por el que se aprueba el texto refundido de la Ley de Propiedad Intelectual."
    ),
    (
        "Queda prohibida su reproducción, distribución, comunicación pública, "
        "modificación, venta, cesión, publicación o difusión, total o parcial, "
        "sin autorización previa y por escrito de su titular."
    ),
    (
        "No está permitida su publicación en páginas web, plataformas digitales, "
        "redes sociales, grupos de mensajería, aplicaciones, repositorios o "
        "servicios de intercambio de archivos."
    ),
]

# A4 width 210 mm minus two margins of 25.4 mm.
CONTENT_WIDTH_DXA = 9026
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "17365D"
BLUE = "1F4E79"
MID_BLUE = "2F75B5"
PALE_BLUE = "D9EAF7"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"
DARK_GRAY = "333333"
WHITE = "FFFFFF"


def set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int]):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, **TABLE_CELL_MARGIN_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bottom_border(paragraph, color: str = "B4C7E7", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=7, color=MID_GRAY)


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1 (\2)", text)
    text = text.replace("`", "")
    return text


def add_inline_runs(paragraph, text: str, *, default_size: float = 11):
    text = clean_inline_markdown(text)
    pieces = re.split(r"(\*\*.+?\*\*|\*[^*]+\*)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif piece.startswith("*") and piece.endswith("*"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, size=default_size, italic=True)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, size=default_size)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def configure_document(doc: Document):
    doc.core_properties.author = AUTHOR
    doc.core_properties.last_modified_by = AUTHOR
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (15, BLUE, 16, 8),
        "Heading 2": (12.5, MID_BLUE, 12, 6),
        "Heading 3": (11.5, BLUE, 8, 4),
        "Heading 4": (11, DARK_GRAY, 6, 3),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.paragraph_format.left_indent = Cm(0.95)
        style.paragraph_format.first_line_indent = Cm(-0.48)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.12

    if "Question" not in styles:
        q_style = styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    else:
        q_style = styles["Question"]
    q_style.font.name = "Arial"
    q_style.font.size = Pt(11)
    q_style.font.bold = True
    q_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    q_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    q_style.paragraph_format.space_before = Pt(10)
    q_style.paragraph_format.space_after = Pt(4)
    q_style.paragraph_format.keep_with_next = True


def configure_headers_and_footers(doc: Document, running_label: str):
    section = doc.sections[0]

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(running_label)
    set_run_font(run, size=8, bold=True, color=BLUE)
    add_bottom_border(p, color="B4C7E7", size="6")

    first_header = section.first_page_header
    first_p = first_header.paragraphs[0]
    first_p.text = ""

    footer = section.footer
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(1)
    compact = (
        f"{COPYRIGHT_LINE} Material protegido por el Real Decreto Legislativo 1/1996. "
        "Prohibida su reproducción, distribución, comunicación pública, modificación, "
        "venta, cesión, publicación o difusión total o parcial sin autorización escrita. "
        "No se permite su publicación en webs, plataformas, redes, mensajería, "
        "aplicaciones, repositorios ni servicios de intercambio."
    )
    run = p1.add_run(compact)
    set_run_font(run, size=6.5, color=MID_GRAY)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("Página ")
    set_run_font(r, size=7, color=MID_GRAY)
    add_page_field(p2)

    first_footer = section.first_page_footer
    fp = first_footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(COPYRIGHT_LINE)
    set_run_font(run, size=7, color=MID_GRAY)


def add_cover(doc: Document, *, kind: str, number: int | None, title: str, block: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(block.upper())
    set_run_font(run, size=9, bold=True, color=MID_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    if kind == "tema":
        label = f"TEMA {number}"
    elif kind == "test":
        label = f"TEST · TEMA {number}"
    else:
        label = kind.upper()
    run = p.add_run(label)
    set_run_font(run, size=13, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=20, bold=True, color=NAVY)
    add_bottom_border(p, color=MID_BLUE, size="12")

    metadata = [
        ("Cuerpo", "Cuerpo Administrativo de la Administración de la Seguridad Social"),
        ("Subgrupo y acceso", "C1 · sistema general de acceso libre"),
        ("Actualización", UPDATE_DATE_TEXT),
        ("Programa", "Anexo I de la Resolución de 22 de diciembre de 2025"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=9.5, bold=True, color=DARK_GRAY)
        r2 = p.add_run(value)
        set_run_font(r2, size=9.5, color=DARK_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("AVISO DE PROPIEDAD INTELECTUAL")
    set_run_font(r, size=8.5, bold=True, color=BLUE)

    for idx, legal in enumerate(LEGAL_PARAGRAPHS):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.right_indent = Cm(0.35)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        add_paragraph_shading(p, LIGHT_GRAY)
        r = p.add_run(legal)
        set_run_font(r, size=8, bold=(idx == 0), color=DARK_GRAY)

    doc.add_page_break()


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", stripped)]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def column_widths(rows: list[list[str]]) -> list[int]:
    col_count = max(len(row) for row in rows)
    weights = []
    for idx in range(col_count):
        lengths = [len(re.sub(r"\W+", "", row[idx])) if idx < len(row) else 1 for row in rows]
        weight = max(8, min(max(lengths), 60))
        weights.append(weight)
    total_weight = sum(weights)
    widths = [max(900, round(CONTENT_WIDTH_DXA * w / total_weight)) for w in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 900:
        deficit = 900 - widths[-1]
        widths[-1] = 900
        widest = max(range(len(widths) - 1), key=lambda i: widths[i])
        widths[widest] -= deficit
    return widths


def add_table(doc: Document, raw_lines: list[str]):
    rows = [split_table_row(line) for line in raw_lines if not is_table_separator(line)]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    rows = [row + [""] * (col_count - len(row)) for row in rows]
    widths = column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            cell.text = ""
            if ridx == 0:
                shade_cell(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if cidx > 0 and len(value) < 16:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, value, default_size=9)
            for run in p.runs:
                if ridx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_source_line(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    add_inline_runs(p, text, default_size=9)


def add_markdown(doc: Document, markdown: str, *, is_test: bool = False):
    lines = markdown.splitlines()
    idx = 0
    in_sources = False
    current_correct_letter = ""
    current_correct_text = ""
    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("<!--"):
            if is_test:
                verification_match = re.search(
                    r"verificaci[oó]n:\s*(.*?);\s*dificultad:\s*(b[aá]sica|media|alta)",
                    stripped,
                    flags=re.IGNORECASE,
                )
                if verification_match and current_correct_letter:
                    verification = clean_inline_markdown(verification_match.group(1).strip())
                    difficulty = (
                        verification_match.group(2)
                        .lower()
                        .replace("basica", "básica")
                    )
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.35)
                    p.paragraph_format.right_indent = Cm(0.2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(7)
                    p.paragraph_format.keep_together = True
                    add_paragraph_shading(p, LIGHT_GRAY)
                    lead = p.add_run(
                        f"Justificación · dificultad {difficulty.upper()} · "
                        f"respuesta {current_correct_letter.lower()}): "
                    )
                    set_run_font(lead, size=8.5, bold=True, color=BLUE)
                    body = (
                        f"«{current_correct_text}» es correcta porque coincide con "
                        f"{verification}. Las demás alternativas modifican, omiten o "
                        "contradicen el dato jurídico aplicable."
                    )
                    answer_run = p.add_run(body)
                    set_run_font(answer_run, size=8.5, color=DARK_GRAY)
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while idx < len(lines):
                candidate = lines[idx].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                idx += 1
            add_table(doc, table_lines)
            continue

        question_match = re.match(r"^(?:#{1,6}\s+)?(\d+)\.\s+(.+)$", stripped)
        if is_test and question_match:
            current_correct_letter = ""
            current_correct_text = ""
            p = doc.add_paragraph(style="Question")
            # Keep each question, its four alternatives and its justification
            # as one visual block whenever the block fits on a page.
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, f"{question_match.group(1)}. {question_match.group(2)}")
            idx += 1
            continue

        option_match = re.match(
            r"^(?:-\s+)?(\*\*)?([a-d])\)\s+(.+?)(\*\*)?$",
            stripped,
            re.IGNORECASE,
        )
        if is_test and option_match:
            correct = bool(option_match.group(1) and option_match.group(4))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_together = True
            p.paragraph_format.keep_with_next = True
            content = f"{option_match.group(2).lower()}) {option_match.group(3)}"
            r = p.add_run(clean_inline_markdown(content).replace("**", ""))
            set_run_font(r, size=10.5, bold=correct)
            if correct:
                current_correct_letter = option_match.group(2)
                current_correct_text = (
                    clean_inline_markdown(option_match.group(3))
                    .replace("**", "")
                    .strip()
                )
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = clean_inline_markdown(heading_match.group(2)).replace("**", "")
            if level == 1:
                idx += 1
                continue
            style_level = min(level - 1, 4)
            p = doc.add_paragraph(style=f"Heading {style_level}")
            p.paragraph_format.keep_with_next = True
            add_inline_runs(p, heading_text, default_size={1: 15, 2: 12.5, 3: 11.5, 4: 11}[style_level])
            in_sources = "Fuentes oficiales utilizadas" in heading_text
            idx += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            if in_sources:
                add_source_line(doc, text)
            else:
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, text)
            idx += 1
            continue

        numbered_match = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if numbered_match and not is_test:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(
                p,
                f"{numbered_match.group(1)}. {numbered_match.group(2)}",
            )
            idx += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(5)
            add_paragraph_shading(p, LIGHT_BLUE)
            add_inline_runs(p, text, default_size=10)
            idx += 1
            continue

        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith("<!--")
                or candidate.startswith("|")
                or re.match(r"^[-*]\s+", candidate)
                or re.match(r"^\d+[.)]\s+", candidate)
                or (is_test and re.match(r"^(\*\*)?[a-d]\)\s+", candidate, re.IGNORECASE))
            ):
                break
            paragraph_lines.append(candidate)
            idx += 1
        text = " ".join(paragraph_lines)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline_runs(p, text)


def safe_filename(text: str, *, max_length: int = 150) -> str:
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r'[<>:"/\\|?*]', " -", text)
    text = re.sub(r"\s+", " ", text).strip(" .")
    if len(text) > max_length:
        text = text[: max_length - 1].rstrip(" .") + "…"
    return text


def source_word_count(markdown: str) -> int:
    without_comments = re.sub(r"<!--.*?-->", "", markdown, flags=re.DOTALL)
    without_urls = re.sub(r"https?://\S+", "", without_comments)
    return len(re.findall(r"\b[\wÁÉÍÓÚÜÑáéíóúüñ]+\b", without_urls, flags=re.UNICODE))


def audit_test(markdown: str) -> dict:
    questions = re.findall(r"(?m)^\s*(?:#{1,6}\s+)?(\d+)\.\s+", markdown)
    correct_options = re.findall(
        r"(?m)^\s*(?:-\s+)?\*\*([a-d])\)\s+.+?\*\*\s*$",
        markdown,
        flags=re.IGNORECASE,
    )
    difficulty_tags = re.findall(r"dificultad:\s*(básica|media|alta)", markdown, flags=re.IGNORECASE)
    option_groups = re.findall(
        r"(?m)^\s*(?:-\s+)?(?:\*\*)?[a-d]\)\s+",
        markdown,
        flags=re.IGNORECASE,
    )
    difficulty_distribution = dict(Counter(x.lower() for x in difficulty_tags))
    return {
        "questions": len(questions),
        "question_numbers": [int(q) for q in questions],
        "correct_marked": len(correct_options),
        "correct_distribution": dict(Counter(letter.lower() for letter in correct_options)),
        "difficulty_distribution": difficulty_distribution,
        "options_total": len(option_groups),
        "valid": (
            len(questions) == 30
            and [int(q) for q in questions] == list(range(1, 31))
            and len(correct_options) == 30
            and len(option_groups) == 120
            and difficulty_distribution == {"básica": 8, "media": 14, "alta": 8}
        ),
    }


def build_theme(theme: dict) -> tuple[Path, dict]:
    source_path = CONTENT_DIR / f"{theme['id']}.md"
    markdown = source_path.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    short_running = (
        f"{theme['id']} · Tema {theme['numero_oficial']} · "
        f"Actualizado {UPDATE_DATE_NUMERIC}"
    )
    configure_headers_and_footers(doc, short_running)
    add_cover(
        doc,
        kind="tema",
        number=theme["numero_oficial"],
        title=theme["titulo"],
        block=theme["bloque"],
    )
    add_markdown(doc, markdown, is_test=False)

    block_dir = "A. Temario general" if theme["id"].startswith("G") else "B. Temario específico Seguridad Social"
    out_dir = ROOT / "02. Temas redactados" / block_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    name = safe_filename(f"Tema {theme['numero_oficial']}. {theme['titulo']}") + ".docx"
    out_path = out_dir / name
    doc.save(out_path)
    return out_path, {
        "source": str(source_path.relative_to(ROOT)),
        "output": str(out_path.relative_to(ROOT)),
        "word_count": source_word_count(markdown),
    }


def build_test(theme: dict) -> tuple[Path, dict]:
    source_path = TEST_DIR / f"{theme['id']}.md"
    markdown = source_path.read_text(encoding="utf-8")
    audit = audit_test(markdown)
    doc = Document()
    configure_document(doc)
    short_running = (
        f"Test {theme['id']} · Tema {theme['numero_oficial']} · "
        f"Actualizado {UPDATE_DATE_NUMERIC}"
    )
    configure_headers_and_footers(doc, short_running)
    add_cover(
        doc,
        kind="test",
        number=theme["numero_oficial"],
        title=theme["titulo"],
        block=theme["bloque"],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "30 preguntas · cuatro alternativas · una sola respuesta correcta, marcada en negrita."
    )
    set_run_font(r, size=9.5, italic=True, color=MID_GRAY)
    add_markdown(doc, markdown, is_test=True)

    block_dir = "A. Temario general" if theme["id"].startswith("G") else "B. Temario específico Seguridad Social"
    out_dir = ROOT / "03. Test por temas" / block_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    name = safe_filename(f"Test Tema {theme['numero_oficial']}. {theme['titulo']}") + ".docx"
    out_path = out_dir / name
    doc.save(out_path)
    audit["source"] = str(source_path.relative_to(ROOT))
    audit["output"] = str(out_path.relative_to(ROOT))
    return out_path, audit


def build_available(selected_ids: set[str] | None = None) -> dict:
    program = json.loads(PROGRAM_PATH.read_text(encoding="utf-8"))
    report = {
        "generated_at": date.today().isoformat(),
        "update_date": UPDATE_DATE_TEXT,
        "themes": {},
        "tests": {},
        "warnings": [],
    }
    for theme in program["temas"]:
        theme_id = theme["id"]
        if selected_ids and theme_id not in selected_ids:
            continue
        theme_source = CONTENT_DIR / f"{theme_id}.md"
        test_source = TEST_DIR / f"{theme_id}.md"
        if theme_source.exists():
            path, info = build_theme(theme)
            report["themes"][theme_id] = info
            if info["word_count"] < 5500:
                report["warnings"].append(
                    f"{theme_id}: tema por debajo de 5.500 palabras ({info['word_count']})."
                )
        if test_source.exists():
            path, audit = build_test(theme)
            report["tests"][theme_id] = audit
            if not audit["valid"]:
                report["warnings"].append(f"{theme_id}: el test no supera la auditoría estructural.")
    QA_DIR.mkdir(parents=True, exist_ok=True)
    report_path = QA_DIR / "informe_generacion.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--ids",
        nargs="*",
        help="Identificadores concretos (por ejemplo G01 G02 S01). Si se omite, construye lo disponible.",
    )
    args = parser.parse_args()
    report = build_available(set(args.ids) if args.ids else None)
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()

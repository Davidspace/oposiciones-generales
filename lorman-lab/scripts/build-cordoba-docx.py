from __future__ import annotations

import hashlib
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SCRIPT = Path(__file__).resolve()
LORMAN_ROOT = SCRIPT.parents[1]
WORKSPACE = SCRIPT.parents[3]
PRIVATE_ROOT = LORMAN_ROOT / "private-content" / "cordoba"
TEMPLATE_ROOT = WORKSPACE / ".codex_tmp" / "cordoba-ss-source-20260812"
OUTPUT_ROOT = PRIVATE_ROOT / "entrega"
REVIEW_DATE = "12 de agosto de 2026"
REVIEW_DATE_SHORT = "12-08-2026"

BLUE = RGBColor(0x2F, 0x75, 0xB5)
DARK_BLUE = RGBColor(0x17, 0x36, 0x5D)
MID_BLUE = RGBColor(0x1F, 0x4E, 0x79)
TEXT = RGBColor(0x33, 0x33, 0x33)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_template(prefix: str) -> Path:
    matches = [path for path in TEMPLATE_ROOT.rglob("*.docx") if path.name.startswith(prefix)]
    if not matches:
        raise FileNotFoundError(f"No se encontró la plantilla {prefix!r} en {TEMPLATE_ROOT}")
    return matches[0]


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_bottom_border(paragraph, color: str = "2F75B5", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_paragraph_shading(paragraph, fill: str = "F1F3F5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_run(paragraph, text: str, *, bold: bool | None = None, size: float | None = None,
            color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_inline(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("`", "")
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            add_run(paragraph, part[2:-2], bold=True)
        else:
            add_run(paragraph, part)


def add_cover(document: Document, *, collection: str, code: str, title: str,
              material: str, version: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, collection.upper(), bold=True, size=9, color=BLUE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, material.upper(), bold=True, size=13, color=MID_BLUE)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, title, bold=True, size=20, color=DARK_BLUE)
    set_paragraph_bottom_border(p)

    for label, value in [
        ("Cuerpo: ", "Auxiliar Administrativo/a del Ayuntamiento de Córdoba"),
        ("Subgrupo y acceso: ", "C2 · sistema general de acceso libre"),
        ("Actualización: ", REVIEW_DATE),
        ("Programa: ", "Anexo II de las bases BOP-A-2024-4049 y rectificación BOP-A-2024-4924"),
        ("Versión: ", version),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, label, bold=True, size=9.5, color=TEXT)
        add_run(p, value, size=9.5, color=TEXT)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    add_run(p, "AVISO DE PROPIEDAD INTELECTUAL", bold=True, size=8.5, color=MID_BLUE)
    for text, bold in [
        ("© 2026 ACADEMIA LORMAN. Todos los derechos reservados.", True),
        ("Material protegido por el Real Decreto Legislativo 1/1996, de 12 de abril, por el que se aprueba el texto refundido de la Ley de Propiedad Intelectual.", False),
        ("Queda prohibida su reproducción, distribución, comunicación pública, modificación, venta, cesión, publicación o difusión total o parcial sin autorización previa y por escrito de su titular.", False),
        ("No está permitida su publicación en páginas web, plataformas digitales, redes sociales, grupos de mensajería, aplicaciones, repositorios o servicios de intercambio de archivos.", False),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=bold, size=8, color=TEXT)
        set_paragraph_shading(p)
        p.paragraph_format.space_after = Pt(0)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def update_header_footer(document: Document, header_text: str) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.clear()
    add_run(header, header_text, bold=True, size=9, color=MID_BLUE)
    set_paragraph_bottom_border(header, size="4")


def add_table(document: Document, rows: list[list[str]]) -> None:
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, row[col_index] if col_index < len(row) else "")
            for run in p.runs:
                run.font.size = Pt(9.5)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = MID_BLUE
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")


def parse_markdown(document: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    skipped_title = False
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue
        if stripped.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if stripped.startswith("| "):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
            parsed = [row for row in parsed if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
            if parsed:
                add_table(document, parsed)
            continue
        if stripped.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
            p.paragraph_format.keep_with_next = True
        elif stripped.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            p.paragraph_format.keep_with_next = True
        elif stripped.startswith("#### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, stripped[5:])
            for run in p.runs:
                run.font.size = Pt(11.5)
            p.paragraph_format.keep_with_next = True
        elif stripped.startswith("> "):
            p = document.add_paragraph()
            add_inline(p, stripped[2:])
            set_paragraph_shading(p)
            p.paragraph_format.space_after = Pt(6)
        elif re.match(r"^[-*] ", stripped):
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = document.add_paragraph()
            add_inline(p, stripped)
            p.paragraph_format.left_indent = Pt(10)
        else:
            p = document.add_paragraph()
            add_inline(p, stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        index += 1


def build_document(template: Path, source: Path, output: Path, *, collection: str,
                   material: str, title: str, version: str, header: str) -> None:
    before = sha256(template)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    document = Document(output)
    clear_body(document)
    update_header_footer(document, header)
    add_cover(document, collection=collection, code=header, title=title, material=material, version=version)
    parse_markdown(document, source)
    document.save(output)
    if sha256(template) != before:
        raise RuntimeError(f"La plantilla original fue modificada: {template}")


def main() -> None:
    theme_template = find_template("Tema 1.")
    simulation_template = find_template("Simulacro 06")

    jobs = [
        (theme_template, PRIVATE_ROOT / "temas" / "Tema 01 - Constitucion Espanola.md", OUTPUT_ROOT / "01. Temas" / "Tema 01 - Constitucion Espanola - Cordoba.docx", "TEMARIO CÓRDOBA", "TEMA 1", "La Constitución Española de 1978: estructura y principios básicos, derechos y deberes fundamentales y organización territorial del Estado", "Tema completo", f"COR01 · Tema 1 · Actualizado {REVIEW_DATE_SHORT}"),
        (theme_template, PRIVATE_ROOT / "temas" / "Tema 06 - Organizacion politica y administrativa I.md", OUTPUT_ROOT / "01. Temas" / "Tema 06 - Organizacion politica y administrativa I - Cordoba.docx", "TEMARIO CÓRDOBA", "TEMA 6", "La organización política y administrativa del Ayuntamiento de Córdoba (I)", "Tema completo · ROM 2025", f"COR06 · Tema 6 · Actualizado {REVIEW_DATE_SHORT}"),
        (simulation_template, PRIVATE_ROOT / "supuestos" / "Supuesto 01 - Alumno - Festival de barrio.md", OUTPUT_ROOT / "02. Supuestos" / "Supuesto 01 - Festival de barrio - Alumno.docx", "AUXILIAR ADMINISTRATIVO/A DEL AYUNTAMIENTO DE CÓRDOBA", "SUPUESTO PRÁCTICO 01", "Festival de barrio y tramitación electrónica", "Versión del alumno · sin soluciones", f"Alumno · Supuesto 01 · {REVIEW_DATE_SHORT}"),
        (simulation_template, PRIVATE_ROOT / "supuestos" / "Supuesto 01 - Resolucion razonada - Festival de barrio.md", OUTPUT_ROOT / "02. Supuestos" / "Supuesto 01 - Festival de barrio - Resolucion razonada.docx", "AUXILIAR ADMINISTRATIVO/A DEL AYUNTAMIENTO DE CÓRDOBA", "SUPUESTO PRÁCTICO 01", "Festival de barrio y tramitación electrónica", "Resolución razonada", f"Profesorado · Supuesto 01 · {REVIEW_DATE_SHORT}"),
        (simulation_template, PRIVATE_ROOT / "supuestos" / "Supuesto 02 - Alumno - Ordenanza fiscal y presupuesto.md", OUTPUT_ROOT / "02. Supuestos" / "Supuesto 02 - Ordenanza fiscal y presupuesto - Alumno.docx", "AUXILIAR ADMINISTRATIVO/A DEL AYUNTAMIENTO DE CÓRDOBA", "SUPUESTO PRÁCTICO 02", "Ordenanza fiscal y presupuesto municipal", "Versión del alumno · sin soluciones", f"Alumno · Supuesto 02 · {REVIEW_DATE_SHORT}"),
        (simulation_template, PRIVATE_ROOT / "supuestos" / "Supuesto 02 - Resolucion razonada - Ordenanza fiscal y presupuesto.md", OUTPUT_ROOT / "02. Supuestos" / "Supuesto 02 - Ordenanza fiscal y presupuesto - Resolucion razonada.docx", "AUXILIAR ADMINISTRATIVO/A DEL AYUNTAMIENTO DE CÓRDOBA", "SUPUESTO PRÁCTICO 02", "Ordenanza fiscal y presupuesto municipal", "Resolución razonada", f"Profesorado · Supuesto 02 · {REVIEW_DATE_SHORT}"),
    ]

    for template, source, output, collection, material, title, version, header in jobs:
        build_document(template, source, output, collection=collection, material=material, title=title, version=version, header=header)
        print(output)


if __name__ == "__main__":
    main()
